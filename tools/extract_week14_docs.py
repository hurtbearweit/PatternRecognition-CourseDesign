from pathlib import Path

from docx import Document


base = Path(r"D:\大三下\模式识别课设\本周进展说明及下周计划报告\十四周")

for path in sorted(base.glob("*.docx")):
    print(f"\n===== FILE: {path.name} =====")
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(text)
    for idx, table in enumerate(doc.tables, start=1):
        print(f"--- TABLE {idx} ---")
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            print(" | ".join(values))
