from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"D:\大三下\模式识别课设\本周进展说明及下周计划报告\十四周")
OUT_FILE = OUT_DIR / "基于改进YOLOv8s的VisDrone目标检测实验报告.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFC7D1"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table)

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = str(text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    return table


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("基于改进 YOLOv8s 的 VisDrone 目标检测实验报告")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Calibri"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(8)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("数据集：VisDroneYOLO    基础模型：YOLOv8s    改进模型：LUD-YOLOv8s")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor.from_string("555555")
    meta.paragraph_format.space_after = Pt(12)

    doc.add_heading("一、实验目的", level=1)
    add_paragraph(
        doc,
        "本实验面向无人机航拍场景中的多目标检测任务，采用 VisDrone 数据集对 YOLOv8s 模型进行训练与改进。"
        "由于 VisDrone 图像中存在目标尺度小、数量密集、遮挡严重以及类别形态差异较大的问题，原始 YOLOv8s "
        "在部分类别上容易出现漏检或定位不准确。因此，本实验在 YOLOv8s 的基础上引入注意力增强模块和多尺度自适应特征融合模块，"
        "以提升模型对复杂航拍场景目标的检测能力。"
    )

    doc.add_heading("二、采用的方法", level=1)
    add_paragraph(
        doc,
        "实验以 YOLOv8s 作为基础检测网络，并参考 LUD-YOLO 的改进思路，对网络结构进行定制化调整。"
        "模型训练使用 YOLO 格式的 VisDrone 数据集，类别数为 10 类，包括 pedestrian、person、bicycle、car、van、truck、"
        "tricycle、awning-tricycle、bus 和 motor。训练过程中使用 960×960 输入尺寸，并结合 Mosaic、Copy-Paste、MixUp "
        "等数据增强方式提高模型的泛化能力。"
    )
    add_bullet(doc, "训练轮数设置为 100 epoch，batch size 设置为 4。")
    add_bullet(doc, "训练设备为 NVIDIA GeForce RTX 4060 Laptop GPU。")
    add_bullet(doc, "训练框架为 Ultralytics YOLO 8.4.53，PyTorch 2.9.1 + CUDA 12.6。")
    add_bullet(doc, "训练完成后使用验证集对 best.pt 权重进行评估。")

    doc.add_heading("三、YOLOv8s 网络结构改进", level=1)
    add_paragraph(
        doc,
        "本实验主要从骨干网络特征提取和检测头特征融合两个方面对原始 YOLOv8s 进行改进。"
    )

    doc.add_heading("1. 引入 C2fBRA 注意力模块", level=2)
    add_paragraph(
        doc,
        "原始 YOLOv8s 中的 C2f 模块具有较好的梯度流动和特征复用能力，但在无人机航拍图像中，复杂背景和小目标密集分布会削弱关键目标特征。"
        "因此，本实验将部分 C2f 模块替换为 C2fBRA 模块，在 Bottleneck 结构中加入注意力增强机制，使网络能够更加关注有效目标区域，"
        "从而增强小目标和遮挡目标的特征表达能力。"
    )

    doc.add_heading("2. 引入 ASFF3 自适应特征融合模块", level=2)
    add_paragraph(
        doc,
        "VisDrone 数据集中目标尺度变化明显，同一图像中可能同时存在车辆、行人、非机动车等多类不同尺度目标。"
        "为增强多尺度特征融合能力，本实验在检测头中引入 ASFF3 模块，对来自不同层级的特征图进行自适应加权融合。"
        "该模块能够根据目标尺度和特征响应动态调整不同尺度特征的贡献，有助于提升模型对多尺度目标的检出能力。"
    )

    add_table(
        doc,
        ["模型", "层数", "参数量", "计算量"],
        [["改进 LUD-YOLOv8s", "173 layers", "17,309,885", "42.1 GFLOPs"]],
        [2500, 1900, 2500, 2460],
    )

    doc.add_heading("四、训练结果", level=1)
    add_paragraph(
        doc,
        "模型训练 100 epoch 后完成验证，验证集共包含 548 张图像和 38,759 个目标实例。整体检测结果如下表所示。"
    )
    add_table(
        doc,
        ["Precision", "Recall", "mAP50", "mAP50-95"],
        [["0.601", "0.486", "0.488", "0.280"]],
        [2340, 2340, 2340, 2340],
    )
    add_paragraph(
        doc,
        "从整体指标来看，模型 Precision 为 0.601，说明预测结果中正确目标占比较高；Recall 为 0.486，说明模型仍存在一定漏检现象。"
        "mAP50 达到 0.488，表明模型在较宽松 IoU 阈值下具备一定检测能力；mAP50-95 为 0.280，说明在更严格定位要求下，"
        "检测框精度仍有进一步提升空间。"
    )

    doc.add_heading("五、分类别检测效果分析", level=1)
    add_table(
        doc,
        ["类别", "Images", "Instances", "P", "R", "mAP50", "mAP50-95"],
        [
            ["pedestrian", "520", "8844", "0.678", "0.528", "0.562", "0.251"],
            ["person", "482", "5125", "0.644", "0.425", "0.442", "0.167"],
            ["bicycle", "364", "1287", "0.428", "0.296", "0.277", "0.123"],
            ["car", "515", "14064", "0.816", "0.828", "0.842", "0.579"],
            ["van", "421", "1975", "0.587", "0.532", "0.523", "0.359"],
            ["truck", "266", "750", "0.595", "0.453", "0.455", "0.296"],
            ["tricycle", "337", "1045", "0.503", "0.415", "0.383", "0.208"],
            ["awning-tricycle", "220", "532", "0.329", "0.239", "0.167", "0.0979"],
            ["bus", "131", "251", "0.779", "0.574", "0.640", "0.456"],
            ["motor", "485", "4886", "0.656", "0.573", "0.585", "0.264"],
        ],
        [1900, 1150, 1350, 1100, 1100, 1300, 1460],
    )
    add_paragraph(
        doc,
        "分类别结果表明，模型对 car 类检测效果最好，mAP50 达到 0.842，mAP50-95 达到 0.579；bus 类检测效果也相对较好。"
        "这说明模型对外观特征明显、尺度较大的车辆类目标具有较强检测能力。相比之下，bicycle、tricycle 和 awning-tricycle "
        "等类别检测效果较差，主要原因是这些类别目标尺度较小、形态变化大，且容易与背景或其他交通目标混淆。"
    )

    doc.add_heading("六、与原始 YOLOv8s 的对比", level=1)
    add_table(
        doc,
        ["模型", "Precision", "Recall", "mAP50", "mAP50-95", "参数量", "GFLOPs"],
        [
            ["原始 YOLOv8s", "0.597", "0.470", "0.476", "0.288", "11.13M", "28.5"],
            ["改进 LUD-YOLOv8s", "0.601", "0.486", "0.488", "0.280", "17.31M", "42.1"],
        ],
        [2100, 1200, 1200, 1200, 1350, 1200, 1110],
    )
    add_paragraph(
        doc,
        "与原始 YOLOv8s 相比，改进模型的 Precision 从 0.597 提升到 0.601，Recall 从 0.470 提升到 0.486，"
        "mAP50 从 0.476 提升到 0.488，说明引入 C2fBRA 和 ASFF3 后，模型对目标的检出能力有一定增强。"
        "但 mAP50-95 从 0.288 下降到 0.280，说明模型在更严格 IoU 阈值下的定位精度略有不足。"
        "同时，改进模型参数量和计算量明显增加，后续仍需要在精度提升和模型复杂度之间进一步平衡。"
    )

    doc.add_heading("七、结论与后续计划", level=1)
    add_paragraph(
        doc,
        "本实验基于 YOLOv8s 构建了改进的 LUD-YOLOv8s 模型，并在 VisDrone 数据集上完成训练与验证。"
        "实验结果表明，改进模型在 Precision、Recall 和 mAP50 指标上相比原始 YOLOv8s 有小幅提升，"
        "说明注意力增强模块和自适应特征融合模块对提升目标检出能力具有一定作用。"
    )
    add_paragraph(
        doc,
        "但从综合指标来看，改进模型的 mAP50-95 略低于原始 YOLOv8s，且模型参数量和计算量增加较多。"
        "因此，后续工作可以从以下方面继续优化：一是调整优化器和学习率策略，使训练参数更充分发挥作用；"
        "二是分别对 C2fBRA 和 ASFF3 进行消融实验，分析各模块的实际贡献；三是进一步优化检测头结构和损失权重，"
        "提高检测框定位精度，尤其是对小目标和遮挡目标的检测效果。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("十四周实验进展报告")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string("777777")

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_document())
