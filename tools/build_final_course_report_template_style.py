from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_final_course_report import (
    ASSET_DIR,
    ROOT,
    copy_representative_images,
    draw_bar_chart,
    make_charts,
    read_text,
)


REPORT_DIR = ROOT / "report"
TEMPLATE = REPORT_DIR / "模式识别与机器学习课程设计报告_可填写模板.docx"
OUT_FILE = REPORT_DIR / "第41组_无人机视角下目标检测_YOLOv8s_IF-YOLO_课程设计报告_第40组风格修订版.docx"


def set_font(run, size=12, bold=None, east_asia="宋体", ascii_font="Times New Roman", color=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_template_styles(doc: Document) -> None:
    """Keep the template's style names, but enforce the requested body typography."""
    for section in doc.sections:
        section.top_margin = Inches(0.89)
        section.bottom_margin = Inches(0.81)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0.56)
        section.footer_distance = Inches(0.68)

    body_style_names = ["Normal", "Body Text", "List Paragraph"]
    for name in body_style_names:
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(12)
        pf = style.paragraph_format
        pf.line_spacing = 1.4875
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Inches(0.29)

    for name, size, east, before, after in [
        ("Heading 1", 18, "Microsoft JhengHei UI", 6.85, 0),
        ("Heading 2", 14, "Microsoft JhengHei UI", 0, 0),
        ("Heading 3", 12, "Microsoft JhengHei UI", 0, 0),
    ]:
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        style.font.size = Pt(size)
        style.font.bold = True
        pf = style.paragraph_format
        pf.line_spacing = 1.15
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True

    if "Code" not in doc.styles:
        doc.styles.add_style("Code", 1)
    code = doc.styles["Code"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    code.font.size = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.first_line_indent = Inches(0)
    code.paragraph_format.space_after = Pt(0)


def paragraph(doc, text="", style="Body Text", align=None, bold=False):
    p = doc.add_paragraph(style=style if style in doc.styles else None)
    p.paragraph_format.line_spacing = 1.4875
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if style in ["Body Text", "Normal"]:
        p.paragraph_format.first_line_indent = Inches(0.29)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, 12, bold=bold)
    return p


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.0
    for r in p.runs:
        if level == 1:
            set_font(r, 18, bold=True, east_asia="Microsoft JhengHei UI", color="000000")
        elif level == 2:
            set_font(r, 14, bold=True, east_asia="Microsoft JhengHei UI", color="000000")
        else:
            set_font(r, 12, bold=True, east_asia="Microsoft JhengHei UI", color="000000")
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, 10.5, east_asia="宋体")
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.4875
    r = p.add_run(text)
    set_font(r, 12)
    return p


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_font(r, size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def table(doc, headers, rows, widths=None, font_size=10.5):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    if widths:
        set_table_width(t, widths)
    for i, header in enumerate(headers):
        shade(t.rows[0].cells[i])
        set_cell(t.rows[0].cells[i], header, bold=True, size=font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 16 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[i], value, align=align, size=font_size)
    doc.add_paragraph()
    return t


def add_image(doc, path, title, width=5.8):
    p = Path(path)
    if not p.exists():
        return False
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.first_line_indent = Inches(0)
    par.paragraph_format.line_spacing = 1.15
    try:
        par.add_run().add_picture(str(p), width=Inches(width))
    except Exception:
        return False
    caption(doc, title)
    return True


def add_image_page(doc, path, title, width=5.75):
    ok = add_image(doc, path, title, width=width)
    if ok:
        doc.add_page_break()
    return ok


def code_block(doc, code, max_lines=80):
    lines = code.splitlines()
    if len(lines) > max_lines:
        lines = lines[:max_lines] + ["# ……以下略，完整文件见项目对应路径。"]
    for line in lines:
        p = doc.add_paragraph(style="Code")
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line[:125])
        set_font(r, 8, east_asia="等线", ascii_font="Consolas")


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_font(run, 9)


def main_chart():
    make_charts()
    return copy_representative_images()


def build_report():
    imgs = main_chart()
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    configure_template_styles(doc)

    # Cover, matching the template's basic layout.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    r = p.add_run("模式识别与机器学习课程")
    set_font(r, 24, bold=True, east_asia="黑体", ascii_font="Times New Roman")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("设计报告")
    set_font(r, 24, bold=True, east_asia="黑体", ascii_font="Times New Roman")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("【基于 YOLOv8s 与 IF-YOLO 改进的无人机视角目标检测】")
    set_font(r, 16, bold=True, east_asia="黑体")
    doc.add_paragraph()
    table(
        doc,
        ["项　　目", "内　　容"],
        [
            ["院　　系", "【填写院系】"],
            ["小组成员", "【填写成员姓名，多人可换行】"],
            ["指导教师", "【填写指导教师】"],
            ["课程班级", "【填写班级】"],
            ["日　　期", "2026 年 6 月"],
        ],
        [2200, 6200],
        font_size=12,
    )
    doc.add_page_break()

    heading(doc, "摘　要", 1)
    for text in [
        "无人机航拍图像具有视角高、目标尺度小、类别密集、遮挡频繁和背景复杂等特点，传统目标检测算法在这类场景中容易出现漏检、误检和定位偏移。本课程设计围绕“无人机视角下的目标检测”展开，以 VisDrone 数据集作为主要实验对象，以 YOLOv8s 为基础算法，完成数据格式转换、统一训练配置筛选、基线复现、多模型横向比较、IF-YOLO 结构改进、消融实验以及多随机种子稳定性测试。",
        "在两个月的实验过程中，我们先比较了输入尺寸、优化器、学习率、数据增强等训练配置，确定后续统一使用的训练方案；随后在同一配置下测试 YOLOv8s、LUD-YOLO、HS-FPN、PRNet、FBRT-YOLO 与 IF-YOLO，最终将 IF-YOLO 作为重点改进对象。在 IF-YOLO 基础上，继续测试 WIoU v3、CSFM-Lite、P5-MKP-Lite、DySample、LocalDetailRefine 和 WeightedFeatureFusion 等模块。",
        "实验结果表明，IF-YOLO 相比原始 YOLOv8s 在召回率与 mAP50 上有明显优势；加入 WIoU v3 与 CSFM-Lite 后，最佳 mAP50-95 提升到 0.30003，较基础 IF-YOLO 提升约 7.3%。最终的 WFF 版本在三随机种子稳定性测试中取得 mAP50-95 均值 0.29380、标准差 0.00138，相比无 WFF 对照组均值 0.28613 更稳定。综合精度、复杂度和稳定性，本报告将 IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3→P2) 作为最终模型。",
    ]:
        paragraph(doc, text)
    paragraph(doc, "关键词：无人机目标检测；YOLOv8s；VisDrone；IF-YOLO；WIoU；CSFM-Lite", bold=True)
    doc.add_page_break()

    heading(doc, "Abstract", 1)
    for text in [
        "Aerial object detection from unmanned aerial vehicles is challenging because targets are usually small, dense, occluded and captured under large scale variation. This course project studies UAV-view object detection on the VisDrone dataset using YOLOv8s as the baseline. The work covers data conversion, unified training setup selection, baseline reproduction, comparison of improved YOLO variants, IF-YOLO-based module design, ablation experiments and multi-seed stability validation.",
        "We first selected a common training configuration, then compared YOLOv8s, LUD-YOLO, HS-FPN, PRNet, FBRT-YOLO and IF-YOLO under the same setting. IF-YOLO was chosen as the main model because it achieved the strongest recall and mAP50 in the model-selection stage. Further experiments tested WIoU v3, CSFM-Lite, P5-MKP-Lite, DySample, LocalDetailRefine and WeightedFeatureFusion modules on top of IF-YOLO.",
        "The best single-run result was obtained by IF-YOLO + WIoU v3 + CSFM-Lite, whose best mAP50-95 reached 0.30003. The final WFF variant showed better stability in three-seed experiments, reaching a mean mAP50-95 of 0.29380 with a standard deviation of 0.00138. Considering accuracy, complexity and stability, IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3→P2) is selected as the final model.",
    ]:
        paragraph(doc, text)
    paragraph(doc, "Key Words: UAV object detection; YOLOv8s; VisDrone; IF-YOLO; WIoU; CSFM-Lite", bold=True)
    doc.add_page_break()

    heading(doc, "目　录", 1)
    toc_items = [
        "摘　要",
        "Abstract",
        "1　项目成员",
        "　1.1　具体工作说明",
        "2　课题概述",
        "　2.1　课题背景与任务内容",
        "　　2.1.1　任务具体要求",
        "　2.2　国内外研究现状",
        "　　2.2.1　相关数据集与公开资源",
        "　　2.2.2　相关方法",
        "　2.3　研究意义与应用场景",
        "3　总体方案与方法设计",
        "　3.1　总体技术路线",
        "　3.2　数据预处理与增强",
        "　3.3　基线模型/基础算法",
        "　3.4　改进模型与关键模块",
        "　　3.4.1　第16周模型拆解与消融",
        "　　3.4.2　横向改进模型",
        "　　3.4.3　最终模型进一步改进",
        "　3.5　损失函数与优化策略",
        "　3.6　算法流程与复杂度分析",
        "4　实验设计与结果分析",
        "　4.1　实验环境与统一训练条件",
        "　4.2　数据集划分与统计",
        "　4.3　评价指标",
        "　4.4　对比实验",
        "　4.5　消融实验",
        "　4.6　可视化结果与误差分析",
        "5　系统实现与应用展示",
        "　5.1　系统总体架构",
        "　5.2　功能模块设计",
        "　5.3　系统运行流程",
        "　5.4　应用效果展示",
        "6　结束语",
        "　6.1　工作总结",
        "　6.2　不足与展望",
        "参考文献",
        "附录 A　核心代码与配置文件",
    ]
    for item in toc_items:
        paragraph(doc, item, style="Normal")
    doc.add_page_break()

    heading(doc, "1　项目成员", 1)
    caption(doc, "表 1-1　项目成员及分工")
    table(
        doc,
        ["姓名", "学号", "年级", "所在院系", "所在班级", "主要分工", "贡献比"],
        [
            ["【成员1】", "【学号】", "2023", "【院系】", "【班级】", "统一配置、YOLOv8s 基线、数据整理", "【%】"],
            ["【成员2】", "【学号】", "2023", "【院系】", "【班级】", "改进模型复现与横向比较", "【%】"],
            ["【成员3】", "【学号】", "2023", "【院系】", "【班级】", "IF-YOLO 消融、WIoU 与 CSFM-Lite", "【%】"],
            ["【成员4】", "【学号】", "2023", "【院系】", "【班级】", "WFF 稳定性、可视化、报告整理", "【%】"],
        ],
        [850, 1050, 650, 1150, 1150, 3250, 720],
        font_size=9,
    )
    heading(doc, "1.1　具体工作说明", 2)
    for item in [
        "完成 VisDrone 原始标注到 YOLO 格式的转换脚本整理，并对数据目录、训练集、验证集和测试集路径进行统一说明。",
        "在第 13 周测试不同输入尺寸、学习率与优化器组合，最后选择 960 输入尺寸、AdamW、lr0=0.003 等统一训练设置。",
        "在统一配置下完成 YOLOv8s、LUD-YOLO、HS-FPN、PRNet、FBRT-YOLO 与 IF-YOLO 的模型训练和横向比较。",
        "围绕 IF-YOLO 完成 WIoU v3、CSFM-Lite、P5-MKP-Lite、DySample、LDR 与 WFF 等模块的实现、安装脚本、训练脚本与结果分析。",
        "完成多随机种子稳定性实验与 1280 输入尺寸复杂度检查，整理实验曲线、混淆矩阵、预测可视化和核心代码附录。",
    ]:
        bullet(doc, item)

    heading(doc, "2　课题概述", 1)
    heading(doc, "2.1　课题背景与任务内容", 2)
    paragraph(doc, "无人机平台能够以较低成本获取大范围、俯视角和动态变化的场景图像，因此在交通监控、城市治理、灾害救援、安防巡检和低空经济应用中具有重要价值。与普通地面视角相比，无人机图像中的目标往往面积小、数量多、分布密集，并受到拍摄高度、俯仰角、光照、运动模糊和遮挡影响，检测难度明显增大。")
    paragraph(doc, "本课题的输入是无人机航拍图像，输出是图像中十类常见目标的位置框、类别和置信度。我们使用 YOLOv8s 作为基础检测器，围绕 VisDrone 数据集开展训练和验证，目标是在尽量保持模型可部署性的前提下提升小目标和密集目标检测能力。")
    add_image(doc, imgs.get("sample_1", ""), "图 2-1　无人机视角下的目标检测任务示意图", width=5.4)
    heading(doc, "2.1.1　任务具体要求", 3)
    for item in [
        "输入为 VisDrone 航拍图像，输出为目标检测框、类别标签和置信度，任务类型为多类别目标检测。",
        "完成 YOLOv8s 基线训练，并在统一训练参数下比较不同改进模型的性能。",
        "围绕最终模型完成消融实验、稳定性测试、可视化展示和关键代码整理。",
    ]:
        bullet(doc, item)
    heading(doc, "2.2　国内外研究现状", 2)
    paragraph(doc, "目标检测方法经历了从手工特征与滑动窗口，到两阶段深度检测器，再到单阶段实时检测器的发展。Faster R-CNN 等两阶段方法精度较高，但推理速度和工程部署成本较高；SSD、YOLO 系列等单阶段方法将候选框生成与类别回归合并到一个网络中，更适合实时检测与边缘部署。")
    paragraph(doc, "在无人机视角目标检测中，研究重点通常集中在三类问题：一是小目标特征在下采样过程中容易丢失，需要更强的浅层信息保持机制；二是不同尺度目标同时出现，需要更合理的多尺度特征融合；三是密集目标定位要求更高，需要改进边界框回归损失或后处理策略。")
    heading(doc, "2.2.1　相关数据集与公开资源", 3)
    table(
        doc,
        ["数据集/资源", "内容", "本项目使用方式"],
        [
            ["VisDrone2019", "无人机航拍目标检测数据集，包含行人、车辆、摩托车等目标", "作为主要训练、验证与测试数据来源"],
            ["Ultralytics YOLOv8", "成熟的目标检测训练与推理框架", "作为 YOLOv8s 基线与自定义模型实现基础"],
            ["项目 Reference 论文", "IF-YOLO、LUD-YOLO、FBRT-YOLO 等参考论文", "用于设计模型对比与模块改进方向"],
        ],
        [1800, 3600, 3600],
        font_size=10,
    )
    caption(doc, "表 2-1　相关数据集与公开资源")
    heading(doc, "2.2.2　相关方法", 3)
    for item in [
        "YOLOv8s：速度快、工程成熟，是本项目的基础模型，但在小目标密集场景中存在漏检和定位不稳问题。",
        "LUD-YOLO：强调轻量化和细节增强，在严格定位指标上表现较好，可作为改进模型对照。",
        "IF-YOLO：通过 IPFA、CSFM 与 FGAFPN 增强小目标信息保持和多尺度融合，是本项目最终选择的主线模型。",
        "WIoU、CSFM-Lite 与 WFF：分别从定位损失、轻量多尺度融合和高分辨率特征加权融合三个角度进一步改进 IF-YOLO。",
    ]:
        bullet(doc, item)
    heading(doc, "2.3　研究意义与应用场景", 2)
    paragraph(doc, "本课题的意义主要体现在三个方面。第一，从算法角度看，它检验了 YOLOv8s 在无人机小目标场景中的瓶颈，并通过消融实验分析不同模块的实际贡献；第二，从工程角度看，项目形成了可复现实验脚本、统一配置、模型检查脚本和结果汇总脚本；第三，从应用角度看，检测结果可为交通流量统计、异常事件发现、目标跟踪与区域巡检提供基础感知能力。")
    add_image(doc, imgs.get("sample_2", ""), "图 2-2　无人机目标检测应用场景示例", width=5.4)

    heading(doc, "3　总体方案与方法设计", 1)
    paragraph(doc, "本章说明从输入数据到输出结果的完整技术路线，包括数据准备、训练配置筛选、基线模型、改进模型、损失函数与复杂度分析。")
    heading(doc, "3.1　总体技术路线", 2)
    paragraph(doc, "本项目的整体路线可以概括为“数据准备、统一训练设置、模型横向筛选、IF-YOLO 深入改进、稳定性验证、固定图像展示”。我们没有一开始就直接堆叠模块，而是先用统一训练参数把不同模型放在同一比较框架下，再选择表现最稳的 IF-YOLO 继续迭代。")
    table(
        doc,
        ["阶段", "处理内容", "输出"],
        [
            ["数据准备", "标注格式转换、路径整理、类别映射", "YOLO 格式数据集"],
            ["统一配置筛选", "比较输入尺寸、学习率、优化器和增强策略", "后续统一训练配置"],
            ["模型横向比较", "同数据、同参数训练多个 YOLO 改进模型", "确定 IF-YOLO 为主线"],
            ["模块改进", "测试 WIoU、CSFM-Lite、MKP、DySample、WFF 等", "最终候选模型"],
            ["稳定性验证", "seed=0/1/2 多次训练", "均值、标准差和最终推荐"],
        ],
        [1600, 4800, 2800],
        font_size=10,
    )
    caption(doc, "图 3-1　总体技术路线图")
    heading(doc, "3.2　数据预处理与增强", 2)
    paragraph(doc, "VisDrone 原始标注以左上角坐标、宽高、类别等字段保存。YOLO 训练要求每个目标表示为归一化后的中心点坐标和宽高，因此我们编写转换脚本，跳过 ignored regions，将类别编号从 VisDrone 的 1 起始调整为 YOLO 的 0 起始。")
    table(
        doc,
        ["配置项", "取值", "说明"],
        [
            ["输入尺寸", "960×960", "第 13 周实验中比 640/800 更适合本数据集"],
            ["batch", "4", "兼顾显存与训练稳定性"],
            ["mosaic", "1.0", "增强密集目标场景泛化能力"],
            ["copy_paste", "0.2", "增加目标组合变化"],
            ["mixup", "0.1", "缓解过拟合并扩展样本分布"],
            ["degrees / scale", "5.0 / 0.3", "模拟拍摄角度与尺度变化"],
            ["seed", "0", "单次对比实验固定随机性"],
        ],
        [1800, 1800, 5600],
        font_size=10,
    )
    caption(doc, "表 3-1　数据预处理配置")
    add_image(doc, imgs.get("fig_labels", ""), "图 3-2　训练数据标签分布可视化", width=5.2)
    heading(doc, "3.3　基线模型/基础算法", 2)
    paragraph(doc, "YOLOv8s 是本课程设计的基础网络，具有较好的速度、参数量和成熟工具链。其骨干网络负责多层特征提取，颈部结构完成多尺度特征融合，检测头输出边界框、类别置信度和目标置信度。对于无人机视角任务，YOLOv8s 的优势是训练和部署成本低，劣势是小目标在连续下采样后容易丢失细节。")
    paragraph(doc, "训练目标可以写成定位损失、分类损失和分布式焦点损失的加权组合：")
    paragraph(doc, "L = λbox Lbox + λcls Lcls + λdfl Ldfl", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc, "式中，Lbox 约束预测框与真实框的几何重合，Lcls 约束类别判断，Ldfl 用于提升边界框回归精度。本项目统一设置 box=7.5、cls=2.0，并保留 YOLOv8 默认的 DFL 机制。")
    heading(doc, "3.4　改进模型与关键模块", 2)
    paragraph(doc, "模型改进部分不是一次性确定的。第 14 周和第 15 周，我们先查阅并复现了 LUD-YOLO、PRNet-YOLO、FBRT-YOLO、HS-FPN、IF-YOLO 等不同方向的改进模型；第 16 周则把重点放到 IF-YOLO 的结构拆解上，分别考察 IPFA、CSFM 和 FGAFPN 对 VisDrone 小目标检测的贡献。完成这一步之后，我们才继续在 IF-YOLO 主线上加入 WIoU v3、CSFM-Lite 和 WFF。这样的顺序能避免只看最终指标，而忽略每个模块实际解决的问题。")
    heading(doc, "3.4.1　第16周模型拆解与消融", 3)
    paragraph(doc, "第 16 周的核心工作是把完整 IF-YOLO 拆成几个可以单独验证的结构。原始 YOLOv8s 的下采样主要依赖步长为 2 的卷积，这种做法速度快，但会在无人机视角中损失小目标的边缘和位置信息。IPFA 模块的改动点是在部分下采样位置用像素重排和通道融合代替普通卷积下采样，将空间邻域信息转移到通道维度，再进行特征提取，目标是在降低分辨率时尽量保留小目标细节。")
    paragraph(doc, "CSFM 模块主要解决多尺度特征融合中的冲突问题。无人机图像中同一张图可能同时存在大车、小行人和摩托车，浅层特征包含细节但语义弱，深层特征语义强但空间信息粗。CSFM 将浅层、中层、深层特征对齐后进行融合，用于增强不同尺度目标之间的互补信息。FGAFPN 则是在 CSFM 的基础上构建更细粒度的四尺度特征金字塔，它不只做传统 P3/P4/P5 三尺度检测，而是引入 P2 高分辨率分支，并在自顶向下和自底向上的路径中重复融合浅层细节，从而更适合密集小目标。")
    table(
        doc,
        ["第16周模型", "具体结构修改", "设计目的", "实验效果"],
        [
            ["YOLOv8s + IPFA", "将部分 Conv 下采样替换为 IPFA，下采样时保留局部空间信息", "减少小目标细节在早期下采样中的丢失", "稳定重训 mAP50-95 为 0.25736，单独使用效果不理想，说明仅保留下采样信息不足以支撑整体提升"],
            ["YOLOv8s + CSFM", "保留普通下采样，在 Head 中加入三组 CSFM 多尺度融合", "让浅层细节、中层纹理和深层语义互补", "mAP50-95 达到 0.29193，明显高于基础 IF-YOLO 的 0.27970，说明多尺度融合贡献较大"],
            ["YOLOv8s + FGAFPN", "使用 CSFM 与四尺度细粒度特征金字塔，引入 P2 高分辨率检测分支", "提升密集小目标和高分辨率目标的检测能力", "mAP50-95 达到 0.30128，是第 16 周消融中表现最好的结构之一"],
            ["完整 IF-YOLO", "组合 IPFA、CSFM 与 FGAFPN，形成四尺度检测输出", "同时保留小目标细节并增强多尺度融合", "第 15 周横向比较中 mAP50 为 0.5103、Recall 为 0.4999，综合表现优于其他候选模型"],
        ],
        [1500, 3000, 2550, 3150],
        font_size=8.5,
    )
    caption(doc, "表 3-3　第16周 IF-YOLO 结构拆解与消融设计")
    heading(doc, "3.4.2　横向改进模型", 3)
    paragraph(doc, "在确定 IF-YOLO 之前，我们还测试了几类不同思路的改进模型。LUD-YOLO 的重点在轻量化和细节增强，主要希望在不大幅增加推理负担的情况下改善小目标定位，因此它在 mAP50-95 上表现较好；PRNet-YOLO 则更强调高分辨率浅层特征的反复利用，通过 ESSamp 在下采样时保留细节，并通过 PRN 结构在特征融合阶段多次引入 P2、P3 层信息；FBRT-YOLO 关注航拍小目标中语义空间不匹配和感受野不足问题，通过 FCM 做语义与空间互补映射，通过 MKP 用多核深度可分离卷积扩大感受野。")
    paragraph(doc, "这些模型的共同目标都是改善无人机小目标检测，但侧重点不同。LUD-YOLO 更偏向定位质量，PRNet-YOLO 更偏向浅层细节复用，FBRT-YOLO 更偏向轻量化多尺度感受野。统一训练后，IF-YOLO 的 Recall 和 mAP50 最好，说明它更适合本项目对“尽量多检出目标”的需求；PRNet-YOLO 虽然思路明确，但训练成本较高，最终指标明显低于其他模型。")
    table(
        doc,
        ["模型", "主要改进点", "优势观察", "不足观察"],
        [
            ["LUD-YOLO", "围绕轻量化上采样和细节增强改进 YOLOv8s", "mAP50-95 为 0.2797，严格定位质量较好", "Recall 和 mAP50 不及 IF-YOLO"],
            ["PRNet-YOLO", "ESSamp 保留下采样细节，PRN 重复利用 P2/P3 高分辨率浅层特征", "结构针对小目标问题较明确", "mAP50-95 仅 0.2287，训练成本与显存压力较高"],
            ["FBRT-YOLO", "FCM 做语义/空间互补映射，MKP 用多核感知扩大感受野", "Precision 较高，训练速度较快", "mAP50-95 为 0.2636，综合效果未超过基线和 IF-YOLO"],
            ["IF-YOLO", "IPFA + CSFM + FGAFPN，四尺度检测输出", "Recall 0.4999、mAP50 0.5103，在横向比较中综合最好", "原始 CSFM 和四尺度结构带来一定计算量"],
        ],
        [1500, 3500, 2400, 2800],
        font_size=8.5,
    )
    caption(doc, "表 3-4　横向改进模型的结构特点与实验观察")
    heading(doc, "3.4.3　最终模型进一步改进", 3)
    paragraph(doc, "在完整 IF-YOLO 基础上，我们继续加入 WIoU v3 和 CSFM-Lite。WIoU v3 的引入目标是提升边界框定位质量，它通过动态聚焦机制调整样本权重，使训练更关注中等质量、仍有改进空间的边界框。实验中，WIoU 对 mAP50 影响不大，但明显提升 mAP50-95，说明它主要改善严格 IoU 阈值下的定位质量。")
    paragraph(doc, "CSFM-Lite 则针对原始 CSFM 参数量和计算量较大的问题，将较重的冲突抑制和空间融合结构替换为全局池化、ECA 通道加权与深度可分离空间分支，在保持多尺度融合思想的同时显著降低模型规模。第 17 周又进一步在 P3→P2 高分辨率融合节点加入 WeightedFeatureFusion，让网络学习不同输入特征的融合权重，而不是简单拼接。WFF 只增加 4098 个参数和约 0.5 GFLOPs，但在三随机种子实验中表现出更高的 mAP50-95 均值和更小波动。")
    add_image(doc, ASSET_DIR / "main_model_comparison.png", "图 3-3　改进模块效果对比图", width=5.8)
    heading(doc, "3.5　损失函数与优化策略", 2)
    paragraph(doc, "优化器统一采用 AdamW，初始学习率 lr0=0.003，最终学习率比例 lrf=0.01，patience=20。相比直接使用默认参数，统一配置使不同模型之间的比较更公平，也减少了“某个模型只是因为参数更合适而获胜”的可能。")
    heading(doc, "3.6　算法流程与复杂度分析", 2)
    table(
        doc,
        ["模型", "Params", "GFLOPs@960", "GFLOPs@1280", "说明"],
        [
            ["YOLOv8s", "11.139M", "约 64.5", "114.7", "基础结构，速度快"],
            ["Basic IF-YOLO", "39.438M", "约 161.5", "287.2", "原始 CSFM 较重"],
            ["IF + WIoU + CSFM-Lite", "11.291M", "105.9", "188.3", "轻量化主干候选"],
            ["Final WFF", "11.295M", "106.4", "189.1", "最终稳定性候选"],
        ],
        [2200, 1500, 1500, 1500, 2700],
        font_size=10,
    )
    caption(doc, "表 3-5　模型复杂度分析")

    heading(doc, "4　实验设计与结果分析", 1)
    heading(doc, "4.1　实验环境与统一训练条件", 2)
    paragraph(doc, "实验主要使用 Ultralytics YOLO 框架完成，训练脚本统一写入项目目录，避免在命令行中随意改动关键超参数。部分实验在服务器环境运行，部分结构检查在本地或 course_ai 环境完成。由于不同阶段硬件可能不同，训练时长仅作为记录，不作为结构优劣判断依据。")
    table(
        doc,
        ["项目", "配置"],
        [
            ["训练框架", "Ultralytics YOLOv8"],
            ["基础权重", "yolov8s.pt"],
            ["训练轮数", "100 epochs"],
            ["输入尺寸", "主要对比 960；第 18 周补充 1280 复杂度与展示包"],
            ["优化器", "AdamW"],
            ["学习率", "lr0=0.003, lrf=0.01"],
            ["评价指标", "Precision、Recall、mAP50、mAP50-95、Params、GFLOPs、FPS"],
        ],
        [2200, 7000],
        font_size=10.5,
    )
    heading(doc, "4.2　数据集划分与统计", 2)
    paragraph(doc, "项目使用 VisDrone2019 检测数据，并转换为 YOLO 格式。训练、验证和测试目录均按 images/ 与 labels/ 分开存放，数据集配置文件只需要指向包含 images 和 labels 的根目录。由于数据集体积较大，完整图像、训练权重和 runs 输出不纳入 Git 仓库，只在本地和服务器保留。")
    heading(doc, "4.3　评价指标", 2)
    paragraph(doc, "Precision 表示预测为正的目标中有多少是真正目标；Recall 表示真实目标中有多少被模型检出；mAP50 是 IoU 阈值为 0.5 时的平均精度；mAP50-95 是 IoU 从 0.5 到 0.95 多阈值平均后的结果，对定位质量要求更严格。本项目将 mAP50-95 作为综合指标，同时关注 Recall 和 mAP50，因为无人机场景中漏检会直接影响应用价值。")
    heading(doc, "4.4　对比实验", 2)
    table(
        doc,
        ["模型", "最佳轮次", "Precision", "Recall", "mAP50", "mAP50-95", "结论"],
        [
            ["YOLOv8s", 89, "0.5990", "0.4841", "0.4856", "0.2670", "基础基线"],
            ["LUD-YOLO", 70, "0.5915", "0.4865", "0.4847", "0.2797", "严格定位较好"],
            ["HS-FPN", 57, "0.5842", "0.4787", "0.4791", "0.2613", "未超过基线"],
            ["PRNet", 85, "0.5454", "0.4285", "0.4245", "0.2287", "训练成本较高且效果弱"],
            ["FBRT-YOLO", 85, "0.5995", "0.4814", "0.4863", "0.2636", "Precision 略好但综合不足"],
            ["IF-YOLO", 89, "0.5965", "0.4999", "0.5103", "0.2797", "综合表现最好，选为主线"],
        ],
        [1500, 900, 1000, 1000, 1000, 1100, 2700],
        font_size=9,
    )
    caption(doc, "表 4-1　不同模型对比实验")
    paragraph(doc, "可以看到，IF-YOLO 的 mAP50 达到 0.5103，Recall 达到 0.4999，是六种模型中最能检出目标的方案；LUD-YOLO 的 mAP50-95 与 IF-YOLO 几乎持平，但 mAP50 和 Recall 低于 IF-YOLO。PRNet-YOLO 虽然通过 ESSamp 和 PRN 反复利用高分辨率细节，但在统一配置下没有取得预期收益，mAP50-95 仅为 0.2287。FBRT-YOLO 的 Precision 较高，但 Recall 与 mAP50-95 不占优势。因此我们没有简单追求单项第一，而是选择综合能力更强、后续改进空间更明确的 IF-YOLO。")
    add_image(doc, imgs.get("fig_yolo_results", ""), "图 4-1　YOLOv8s 基线训练曲线", width=5.6)
    add_image(doc, imgs.get("fig_if_results", ""), "图 4-2　IF-YOLO 训练曲线", width=5.6)
    heading(doc, "4.5　消融实验", 2)
    table(
        doc,
        ["模型/模块", "最佳轮次", "Precision", "Recall", "mAP50", "mAP50-95", "相对基础 IF-YOLO"],
        [
            ["基础 IF-YOLO", 89, "0.59654", "0.49985", "0.51025", "0.27970", "-"],
            ["仅 IPFA", 85, "未记录", "未记录", "0.46986", "0.25736", "低于完整 IF-YOLO"],
            ["仅 CSFM", 90, "未记录", "未记录", "0.51622", "0.29193", "+0.01223"],
            ["FGAFPN", 79, "未记录", "未记录", "0.52070", "0.30128", "+0.02158"],
            ["IF-YOLO + WIoU v3", 84, "0.60004", "0.50065", "0.51042", "0.29014", "+0.01044"],
            ["IF + WIoU + CSFM-Lite", 89, "0.60453", "0.50807", "0.51578", "0.30003", "+0.02033"],
            ["+ P5-MKP-Lite", 88, "0.60144", "0.50032", "0.50919", "0.28740", "+0.00770"],
            ["DySample 双节点", 89, "0.60604", "0.48764", "0.51224", "0.29768", "接近但未超过"],
            ["WFF 单次", 86, "0.60519", "0.48976", "0.51445", "0.29896", "接近 CSFM-Lite"],
        ],
        [1900, 900, 1000, 1000, 1000, 1100, 2300],
        font_size=9,
    )
    caption(doc, "表 4-2　消融实验结果")
    paragraph(doc, "第 16 周消融可以看出，IPFA 单独加入时并没有带来稳定提升，说明小目标检测不能只依赖下采样细节保留，还需要后续多尺度融合共同配合；CSFM 单独加入后 mAP50-95 提升到 0.29193，证明多尺度特征融合是 IF-YOLO 中贡献较大的部分；FGAFPN 进一步利用四尺度和高分辨率 P2 分支，mAP50-95 达到 0.30128，说明细粒度特征金字塔对 VisDrone 这类密集小目标数据集非常重要。")
    paragraph(doc, "在完整 IF-YOLO 基础上，WIoU v3 带来的主要收益体现在 mAP50-95，说明定位质量变好；CSFM-Lite 在保持轻量化的同时进一步提升 Precision、Recall、mAP50 和 mAP50-95，是第 16 周后半阶段最有效的结构改进。P5-MKP-Lite 没有形成叠加收益，DySample 单节点效果较弱，双节点和 WFF 接近 CSFM-Lite，但仍需稳定性验证。")
    add_image(doc, imgs.get("fig_final_results", ""), "图 4-3　IF-YOLO + WIoU v3 + CSFM-Lite 训练曲线", width=5.6)
    add_image(doc, imgs.get("fig_final_pr", ""), "图 4-4　CSFM-Lite 最优模型 PR 曲线", width=5.4)
    add_image(doc, ASSET_DIR / "multiseed_stability.png", "图 4-5　三随机种子稳定性对比", width=5.4)
    heading(doc, "4.6　可视化结果与误差分析", 2)
    paragraph(doc, "从预测图可以看出，模型对车辆、公交车等尺寸较大且纹理清晰的类别检测效果较好；对行人、person、bicycle、awning-tricycle 等小目标和细长目标仍存在漏检。主要原因包括：目标在图像中占比过小；密集区域中目标之间互相遮挡；部分类别外观差异小；YOLO txt 验证流程没有直接输出 COCO 风格 APsmall，因此对极小目标的定量分析还不够细。")
    add_image(doc, imgs.get("fig_final_conf", ""), "图 4-6　CSFM-Lite 归一化混淆矩阵", width=5.2)
    add_image(doc, imgs.get("fig_final_val", ""), "图 4-7　CSFM-Lite 候选模型验证集预测示例", width=5.7)
    add_image(doc, imgs.get("fig_wff_val", ""), "图 4-8　最终 WFF 候选模型验证集预测示例", width=5.7)
    paragraph(doc, "为了让实验过程更容易复核，下面继续给出各阶段代表性曲线与预测图。这些图均来自项目训练输出目录，主要用于说明模型从基线筛选到 IF-YOLO 改进、再到最终稳定性验证的完整过程。")
    gallery = [
        (ROOT / "Model/runs/detect/week15_yolov8s_baseline/BoxPR_curve.png", "图 4-9　YOLOv8s 基线 PR 曲线"),
        (ROOT / "Model/runs/detect/week15_yolov8s_baseline/confusion_matrix_normalized.png", "图 4-10　YOLOv8s 基线归一化混淆矩阵"),
        (ROOT / "Model/runs/detect/week15_lud_yolo/results.png", "图 4-11　LUD-YOLO 训练曲线"),
        (ROOT / "Model/runs/detect/week15_lud_yolo/BoxPR_curve.png", "图 4-12　LUD-YOLO PR 曲线"),
        (ROOT / "本周进展说明及下周计划报告/十五周/train-2/results.png", "图 4-13　HS-FPN 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十五周/train-35/results.png", "图 4-14　PRNet 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十五周/fbrt_yolov8s_visdrone_pretrained_week18-4/results.png", "图 4-15　FBRT-YOLO 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/ablation_03_csfm/ablation_03_csfm/results.png", "图 4-16　CSFM 消融训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/detect/ablation_04_fgafpn/results.png", "图 4-17　FGAFPN 消融训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/IPFA_new/ablation_02_ipfa_stable/ablation_02_ipfa_stable/results.png", "图 4-18　IPFA 稳定版消融训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou-2/week15_if_yolov8s_wiou-2/results.png", "图 4-19　IF-YOLO + WIoU v3 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite_mkp/week15_if_yolov8s_wiou_csfm_lite_mkp/results.png", "图 4-20　P5-MKP-Lite 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_p4p3/results.png", "图 4-21　DySample P4→P3 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_p3p2/results.png", "图 4-22　DySample P3→P2 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_p4p3_p3p2/results.png", "图 4-23　DySample 双节点训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_ldr/results.png", "图 4-24　LocalDetailRefine 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/confusion_matrix_normalized.png", "图 4-25　WFF 归一化混淆矩阵"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/val_batch1_pred.jpg", "图 4-26　WFF 验证集预测示例"),
    ]
    for path, title in gallery:
        add_image_page(doc, path, title)

    heading(doc, "5　系统实现与应用展示", 1)
    heading(doc, "5.1　系统总体架构", 2)
    paragraph(doc, "项目实现并不是单一训练脚本，而是一套可复现实验流程：Dataset 目录负责数据转换；Model 目录负责基线与通用训练脚本；各周报告目录保留阶段性模型包、模块安装脚本、训练脚本、结果汇总脚本和可视化图；Reference 目录保存相关论文；report 目录保存课程报告模板和最终报告。")
    table(
        doc,
        ["目录/文件", "作用"],
        [
            ["Dataset/visdrone2yolo.py", "将 VisDrone 原始标注转换为 YOLO 标签"],
            ["Model/train_visdrone.py", "YOLOv8s 基线训练入口"],
            ["Model/train_config.yaml", "第 15 周后统一训练配置"],
            ["十六周/week15_final_training_package", "IF-YOLO、WIoU、CSFM-Lite、MKP 训练包"],
            ["十七周/week17", "WFF 与多随机种子稳定性实验包"],
            ["18周", "1280 输入尺寸复杂度检查与固定测试图展示包"],
        ],
        [3600, 5600],
        font_size=10,
    )
    heading(doc, "5.2　功能模块设计", 2)
    paragraph(doc, "核心实现模块包括数据转换、模型注册、训练入口、模型检查、结果汇总和固定图片预测。自定义模块需要复制到 Ultralytics 对应模块目录，并在 tasks.py 的模型解析过程中注册通道推导逻辑。为了减少人工改源码出错，项目将这些步骤写成 install_weekXX_modules.py，保证训练环境可重复搭建。")
    heading(doc, "5.3　系统运行流程", 2)
    for item in [
        "准备 VisDroneYOLO 数据集，并在 visdrone.yaml 中填写真实 path。",
        "安装 Python、PyTorch、Ultralytics，并确认 GPU 可用。",
        "运行自定义模块安装脚本，检查 IPFA、CSFM-Lite、WFF 等模块能否构建。",
        "使用 train_experiment.py 或 train_1280_experiment.py 按模型名称启动训练。",
        "训练结束后读取 results.csv、results.png、PR 曲线、混淆矩阵和预测图。",
        "运行 summarize 脚本汇总多随机种子或多模型结果。",
    ]:
        bullet(doc, item)
    heading(doc, "5.4　应用效果展示", 2)
    paragraph(doc, "第 18 周准备了固定 15 张 VisDrone test-challenge 图片，用于让不同模型在完全相同的图像上输出预测结果。该展示包不参与训练和验证指标计算，主要用于课程答辩中直观展示模型对车辆、行人、摩托车等目标的检测效果。")
    add_image(doc, imgs.get("sample_3", ""), "图 5-1　固定测试图片示例", width=5.5)
    add_image(doc, imgs.get("sample_4", ""), "图 5-2　固定测试图片示例", width=5.5)

    heading(doc, "6　结束语", 1)
    heading(doc, "6.1　工作总结", 2)
    paragraph(doc, "本课程设计从 YOLOv8s 基线出发，围绕无人机视角目标检测完成了比较完整的实验闭环。我们没有只跑一个最终模型，而是先确定统一训练配置，再横向比较多个改进方向，最后围绕 IF-YOLO 做进一步模块消融和稳定性验证。最终选择 IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3→P2)，主要原因是它在三随机种子测试中表现更稳定，且复杂度仅比无 WFF 版本略有增加。")
    heading(doc, "6.2　不足与展望", 2)
    paragraph(doc, "本项目仍有不足。第一，虽然 mAP50-95 有提升，但小目标和严重遮挡类别仍然是短板；第二，当前验证流程主要使用 YOLO 自带指标，未进一步启用 COCO-json 评估以获得 APsmall、APmedium 等尺度指标；第三，部分模型训练硬件不同，训练时长只能参考，不能直接比较结构效率；第四，最终模型尚未部署到真实无人机视频流或边缘设备上。后续可以继续从尺度敏感评价、轻量化部署、视频时序信息和更强的数据增强策略四个方向改进。")

    heading(doc, "参考文献", 1)
    for ref in [
        "[1] Ultralytics. YOLOv8 documentation and implementation.",
        "[2] Du D., Zhu P., Wen L., et al. VisDrone-DET2019: The Vision Meets Drone Object Detection in Image Challenge.",
        "[3] IF-YOLO: improved YOLO-based object detection method for drone-view or small-object detection scenarios.",
        "[4] LUD-YOLO: lightweight UAV/drone object detection method.",
        "[5] FBRT-YOLO: feature-enhanced YOLO variant used as a comparative method in this course project.",
        "[6] Zheng Z., Wang P., Liu W., et al. Distance-IoU Loss: Faster and Better Learning for Bounding Box Regression.",
        "[7] 项目本地实验记录：D:\\大三下\\模式识别课设\\findings.md、progress.md 与各周阶段报告。",
    ]:
        paragraph(doc, ref)

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "附录 A　核心代码与配置文件", 1)
    heading(doc, "A.1　项目目录结构", 2)
    code_block(
        doc,
        """.
├─ Dataset/                         数据转换脚本与数据说明
├─ Model/                           YOLOv8s 基线、统一配置和训练脚本
├─ Reference/                       参考论文
├─ report/                          模板与最终报告
├─ tools/                           报告生成、文档提取、结果对比工具
└─ 本周进展说明及下周计划报告/       第 12-18 周阶段材料、实验包和结果""",
        30,
    )
    heading(doc, "A.2　统一训练配置", 2)
    code_block(doc, read_text(ROOT / "Model/train_config.yaml"), 80)
    heading(doc, "A.2.1　VisDrone 数据集配置", 3)
    code_block(doc, read_text(ROOT / "Model/visdrone.yaml"), 80)
    doc.add_page_break()
    heading(doc, "A.3　核心代码", 2)
    paragraph(doc, "以下代码为报告中涉及的关键源代码片段，完整代码可在项目对应路径中查看。")
    heading(doc, "A.3.1　VisDrone 标注转换脚本", 3)
    code_block(doc, read_text(ROOT / "Dataset/visdrone2yolo.py"), 120)
    doc.add_page_break()
    heading(doc, "A.3.2　YOLOv8s 基线训练入口", 3)
    code_block(doc, read_text(ROOT / "Model/train_visdrone.py"), 120)
    doc.add_page_break()
    heading(doc, "A.3.3　第十五周统一训练入口", 3)
    p = ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_package/train_week15_experiment.py"
    if p.exists():
        code_block(doc, read_text(p), 180)
    heading(doc, "A.3.4　IF-YOLO + WIoU + CSFM-Lite YAML", 3)
    p = ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_package/if_yolov8s_wiou_csfm_lite.yaml"
    if p.exists():
        code_block(doc, read_text(p), 180)
    doc.add_page_break()
    heading(doc, "A.3.5　WeightedFeatureFusion 核心模块", 3)
    p = ROOT / "本周进展说明及下周计划报告/十七周/week17/modules/weighted_fusion.py"
    if p.exists():
        code_block(doc, read_text(p), 180)
    heading(doc, "A.3.6　最终 WFF YAML 配置", 3)
    p = ROOT / "本周进展说明及下周计划报告/十七周/week17/configs/if_yolov8s_wiou_csfmlite_wff_p3p2.yaml"
    if p.exists():
        code_block(doc, read_text(p), 180)
    doc.add_page_break()
    heading(doc, "A.3.7　稳定性实验训练入口", 3)
    p = ROOT / "本周进展说明及下周计划报告/十七周/week17/scripts/train_experiment.py"
    if p.exists():
        code_block(doc, read_text(p), 150)
    heading(doc, "A.3.8　第 18 周 1280 输入检查脚本", 3)
    p = ROOT / "本周进展说明及下周计划报告/18周/scripts/check_models_1280.py"
    if p.exists():
        code_block(doc, read_text(p), 95)

    for section in doc.sections:
        add_page_number(section)
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
    doc.core_properties.title = "基于 YOLOv8s 与 IF-YOLO 改进的无人机视角目标检测课程设计报告"
    doc.core_properties.author = "第41组"
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_report())
