from __future__ import annotations

import csv
import json
import re
import shutil
import textwrap
import zipfile
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
TEMPLATE = REPORT_DIR / "模式识别与机器学习课程设计报告_可填写模板.docx"
OUT_FILE = REPORT_DIR / "第41组_无人机视角下目标检测_YOLOv8s_IF-YOLO_课程设计报告.docx"
ASSET_DIR = REPORT_DIR / "final_report_assets"


def set_run_font(run, size=None, bold=None, color=None, font="Calibri", east_asia="宋体"):
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)
        section.header_distance = Inches(0.45)
        section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Inches(0.29)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, "1F4D78", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 5),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code" not in styles:
        styles.add_style("Code", 1)
    code = styles["Code"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    code.font.size = Pt(8)
    code.paragraph_format.first_line_indent = Inches(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run, size=9, color="666666")


def para(doc, text="", style=None, align=None, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    return p


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, east_asia="黑体")
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("• " + text)
    set_run_font(r)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, size=9, color="555555")
    return p


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(w))
            tc_w.set(qn("w:type"), "dxa")
    mar = tbl_pr.find(qn("w:tblCellMar")) or OxmlElement("w:tblCellMar")
    if mar.getparent() is None:
        tbl_pr.append(mar)
    for k, v in {"top": "80", "bottom": "80", "start": "100", "end": "100"}.items():
        node = mar.find(qn(f"w:{k}")) or OxmlElement(f"w:{k}")
        if node.getparent() is None:
            mar.append(node)
        node.set(qn("w:w"), v)
        node.set(qn("w:type"), "dxa")


def table(doc, headers, rows, widths=None, font_size=8.8):
    t = doc.add_table(rows=1, cols=len(headers))
    for style_name in ("Table Grid", "网格型", "Table Normal"):
        try:
            t.style = style_name
            break
        except Exception:
            pass
    if widths:
        set_table_geometry(t, widths)
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], "E8EEF5")
        set_cell(t.rows[0].cells[i], h, bold=True, size=font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(val)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[i], val, align=align, size=font_size)
    doc.add_paragraph()
    return t


def add_image(doc, path, title, width=5.9):
    path = Path(path)
    if not path.exists():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    try:
        p.add_run().add_picture(str(path), width=Inches(width))
    except Exception:
        return False
    caption(doc, title)
    return True


def add_image_plate(doc, path, title, width=6.25):
    ok = add_image(doc, path, title, width=width)
    if ok:
        doc.add_page_break()
    return ok


def add_code_block(doc, code, max_lines=70):
    lines = code.splitlines()
    if len(lines) > max_lines:
        lines = lines[:max_lines] + ["# ……以下代码略，完整文件见项目对应路径。"]
    for line in lines:
        p = doc.add_paragraph(style="Code")
        p.paragraph_format.first_line_indent = Inches(0)
        r = p.add_run(line[:120])
        set_run_font(r, size=7.5, font="Consolas", east_asia="等线")


def read_text(path, limit=None):
    text = Path(path).read_text(encoding="utf-8", errors="ignore")
    return text if limit is None else text[:limit]


def best_row(csv_path):
    df = pd.read_csv(csv_path)
    idx = df["metrics/mAP50-95(B)"].idxmax()
    row = df.loc[idx].to_dict()
    return {
        "epoch": int(row["epoch"]),
        "precision": float(row["metrics/precision(B)"]),
        "recall": float(row["metrics/recall(B)"]),
        "map50": float(row["metrics/mAP50(B)"]),
        "map5095": float(row["metrics/mAP50-95(B)"]),
    }


def rel(path):
    return str(Path(path).relative_to(ROOT))


def make_charts():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    main_models = [
        ("YOLOv8s", ROOT / "Model/runs/detect/week15_yolov8s_baseline/results.csv"),
        ("LUD-YOLO", ROOT / "Model/runs/detect/week15_lud_yolo/results.csv"),
        ("IF-YOLO", ROOT / "本周进展说明及下周计划报告/十五周/train-5/results.csv"),
        ("IF+WIoU", ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou-2/week15_if_yolov8s_wiou-2/results.csv"),
        ("IF+WIoU+CSFM-Lite", ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite-3/week15_if_yolov8s_wiou_csfm_lite-3/results.csv"),
        ("Final WFF", None),
    ]
    rows = []
    for name, p in main_models:
        if p and p.exists():
            b = best_row(p)
            rows.append((name, b["map50"], b["map5095"]))
        elif name == "Final WFF":
            rows.append((name, 0.51180, 0.29380))
    out = ASSET_DIR / "main_model_comparison.png"
    draw_bar_chart(
        out,
        "Main model comparison under unified training settings",
        [r[0] for r in rows],
        [("mAP50", [r[1] for r in rows], "#4C78A8"), ("mAP50-95", [r[2] for r in rows], "#F58518")],
        y_min=0.2,
        y_max=0.55,
    )

    ms = ROOT / "本周进展说明及下周计划报告/十七周结果/tables/tables/multiseed_summary.csv"
    if ms.exists():
        df = pd.read_csv(ms)
        sub = df[df["metric"].isin(["map50", "map50_95"])]
        labels = []
        means_5095 = []
        means_50 = []
        stds_5095 = []
        for i, metric in enumerate(["map50", "map50_95"]):
            for g in ["A", "B"]:
                row = sub[(sub["group"] == g) & (sub["metric"] == metric)].iloc[0]
                if metric == "map50":
                    means_50.append(row["mean"])
                else:
                    labels.append(g)
                    means_5095.append(row["mean"])
                    stds_5095.append(row["std"])
        draw_bar_chart(
            ASSET_DIR / "multiseed_stability.png",
            "Multi-seed stability: A baseline vs B final WFF",
            labels,
            [("mAP50", means_50, "#4C78A8"), ("mAP50-95", means_5095, "#F58518")],
            y_min=0.25,
            y_max=0.53,
            notes=[f"std={s:.4f}" for s in stds_5095],
        )
    return out


def chart_font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size)
        except Exception:
            continue
    return ImageFont.load_default()


def draw_bar_chart(out_path, title, labels, series, y_min=0.0, y_max=1.0, notes=None):
    width, height = 1600, 850
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    title_font = chart_font(34, True)
    label_font = chart_font(20)
    small_font = chart_font(18)
    axis_font = chart_font(18)
    left, right, top, bottom = 110, 60, 95, 170
    plot_w = width - left - right
    plot_h = height - top - bottom
    d.text((left, 25), title, fill="#1F2937", font=title_font)

    for i in range(6):
        yv = y_min + (y_max - y_min) * i / 5
        y = top + plot_h - (yv - y_min) / (y_max - y_min) * plot_h
        d.line((left, y, width - right, y), fill="#E5E7EB", width=2)
        d.text((25, y - 12), f"{yv:.2f}", fill="#4B5563", font=axis_font)
    d.line((left, top, left, top + plot_h), fill="#374151", width=2)
    d.line((left, top + plot_h, width - right, top + plot_h), fill="#374151", width=2)

    group_w = plot_w / max(1, len(labels))
    bar_w = min(70, group_w / (len(series) + 1.4))
    for gi, lab in enumerate(labels):
        cx = left + group_w * gi + group_w / 2
        for si, (_, values, color) in enumerate(series):
            value = values[gi]
            x0 = cx - (len(series) * bar_w) / 2 + si * bar_w
            x1 = x0 + bar_w * 0.82
            y1 = top + plot_h
            y0 = top + plot_h - (value - y_min) / (y_max - y_min) * plot_h
            d.rectangle((x0, y0, x1, y1), fill=color)
            d.text((x0 - 4, y0 - 28), f"{value:.3f}", fill="#111827", font=small_font)
        wrapped = lab.replace(" + ", "\n+ ")
        d.multiline_text((cx - group_w * 0.43, top + plot_h + 18), wrapped, fill="#111827", font=label_font, spacing=2)
        if notes:
            d.text((cx - group_w * 0.25, top + plot_h + 95), notes[gi], fill="#6B7280", font=small_font)

    lx = width - right - 330
    ly = 35
    for name, _, color in series:
        d.rectangle((lx, ly, lx + 24, ly + 24), fill=color)
        d.text((lx + 35, ly - 2), name, fill="#111827", font=label_font)
        ly += 35
    img.save(out_path)


def copy_representative_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image_map = {
        "fig_yolo_results": ROOT / "Model/runs/detect/week15_yolov8s_baseline/results.png",
        "fig_if_results": ROOT / "本周进展说明及下周计划报告/十五周/train-5/results.png",
        "fig_final_results": ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite-3/week15_if_yolov8s_wiou_csfm_lite-3/results.png",
        "fig_final_pr": ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite-3/week15_if_yolov8s_wiou_csfm_lite-3/BoxPR_curve.png",
        "fig_final_conf": ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite-3/week15_if_yolov8s_wiou_csfm_lite-3/confusion_matrix_normalized.png",
        "fig_final_val": ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite-3/week15_if_yolov8s_wiou_csfm_lite-3/val_batch0_pred.jpg",
        "fig_wff_results": ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/results.png",
        "fig_wff_pr": ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/BoxPR_curve.png",
        "fig_wff_val": ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/val_batch0_pred.jpg",
        "fig_labels": ROOT / "Model/runs/detect/week15_yolov8s_baseline/labels.jpg",
    }
    copied = {}
    for key, src in image_map.items():
        if src.exists():
            dst = ASSET_DIR / f"{key}{src.suffix}"
            shutil.copyfile(src, dst)
            copied[key] = dst
    samples = sorted((ROOT / "本周进展说明及下周计划报告/18周/sample_images").glob("*.jpg"))[:4]
    for i, src in enumerate(samples, start=1):
        dst = ASSET_DIR / f"sample_{i}.jpg"
        shutil.copyfile(src, dst)
        copied[f"sample_{i}"] = dst
    return copied


def build_report():
    make_charts()
    imgs = copy_representative_images()

    doc = Document(str(TEMPLATE))
    clear_doc(doc)
    configure_styles(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("模式识别与机器学习课程设计报告")
    set_run_font(r, size=22, bold=True, color="0B2545", east_asia="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 YOLOv8s 与 IF-YOLO 改进的无人机视角目标检测")
    set_run_font(r, size=18, bold=True, color="1F4D78", east_asia="黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VisDrone 数据集 · 统一配置对比 · 模块消融 · 多随机种子稳定性测试")
    set_run_font(r, size=11, color="555555")
    doc.add_paragraph()
    table(
        doc,
        ["项目", "内容"],
        [
            ["课程设计题目", "无人机视角下的目标检测"],
            ["基础算法", "YOLOv8s"],
            ["最终模型", "IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3->P2)"],
            ["数据集", "VisDrone2019 转换为 YOLO 格式，10 类目标"],
            ["小组成员", "【请填写成员姓名、学号、班级】"],
            ["指导教师", "【请填写指导教师】"],
            ["日期", "2026 年 6 月"],
        ],
        [2200, 7160],
        font_size=10,
    )
    doc.add_page_break()

    heading(doc, "摘　要", 1)
    para(doc, "无人机航拍图像具有视角高、目标尺度小、类别密集、遮挡频繁和背景复杂等特点，传统目标检测算法在这类场景中容易出现漏检、误检和定位偏移。本课程设计围绕“无人机视角下的目标检测”展开，选用 VisDrone 数据集作为主要实验对象，以 YOLOv8s 为基础算法，完成数据格式转换、统一训练配置筛选、基线复现、多模型横向比较、IF-YOLO 结构改进、消融实验以及多随机种子稳定性测试。")
    para(doc, "在两个月的实验过程中，我们先比较了输入尺寸、优化器、学习率、数据增强等训练配置，确定了后续统一使用的训练方案；随后在同一配置下对 YOLOv8s、LUD-YOLO、HS-FPN、PRNet、FBRT-YOLO 与 IF-YOLO 进行横向比较，最终将 IF-YOLO 作为重点改进对象。在 IF-YOLO 基础上，我们围绕小目标信息保持、定位损失、多尺度特征融合和高分辨率特征融合继续设计实验，依次测试了 WIoU v3、CSFM-Lite、P5-MKP-Lite、DySample、LocalDetailRefine 和 WeightedFeatureFusion 等模块。")
    para(doc, "实验结果表明，IF-YOLO 相比原始 YOLOv8s 在召回率与 mAP50 上有明显优势；加入 WIoU v3 与 CSFM-Lite 后，最佳 mAP50-95 提升到 0.30003，较基础 IF-YOLO 提升约 7.3%。最终的 WFF 版本在三随机种子稳定性测试中取得 mAP50-95 均值 0.29380、标准差 0.00138，相比无 WFF 的对照组均值 0.28613 更稳定。综合精度、复杂度和稳定性，本报告将 IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3->P2) 作为最终展示模型，同时保留 CSFM-Lite 单模型作为最高单次精度结果。")
    para(doc, "关键词：无人机目标检测；YOLOv8s；VisDrone；IF-YOLO；WIoU；CSFM-Lite", bold=True)
    doc.add_page_break()

    heading(doc, "Abstract", 1)
    para(doc, "Aerial object detection from unmanned aerial vehicles is challenging because objects are usually small, densely distributed, partially occluded, and captured under large scale variation. This course project studies UAV-view object detection on the VisDrone dataset using YOLOv8s as the baseline. The work covers data conversion, unified training setup selection, baseline reproduction, comparison of improved YOLO variants, IF-YOLO-based module design, ablation experiments, and multi-seed stability validation.")
    para(doc, "We first selected a common training configuration, then compared YOLOv8s, LUD-YOLO, HS-FPN, PRNet, FBRT-YOLO, and IF-YOLO under the same setting. IF-YOLO was chosen as the main model because it achieved the strongest recall and mAP50 in the model-selection stage. Further experiments tested WIoU v3, CSFM-Lite, P5-MKP-Lite, DySample, LocalDetailRefine, and WeightedFeatureFusion modules on top of IF-YOLO.")
    para(doc, "The best single-run result was obtained by IF-YOLO + WIoU v3 + CSFM-Lite, whose best mAP50-95 reached 0.30003. The final WFF variant showed better stability in three-seed experiments, reaching a mean mAP50-95 of 0.29380 with a standard deviation of 0.00138. Considering accuracy, complexity and stability, IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3->P2) is selected as the final model for this report.")
    para(doc, "Key Words: UAV object detection; YOLOv8s; VisDrone; IF-YOLO; WIoU; CSFM-Lite", bold=True)
    doc.add_page_break()

    heading(doc, "目　录", 1)
    for item in [
        "1 项目成员",
        "2 课题概述",
        "3 总体方案与方法设计",
        "4 实验设计与结果分析",
        "5 系统实现与应用展示",
        "6 结束语",
        "参考文献",
        "附录 A 核心代码与配置文件",
    ]:
        para(doc, item)
    doc.add_page_break()

    heading(doc, "1　项目成员", 1)
    table(doc, ["姓名", "学号", "年级", "所在院系", "所在班级", "主要分工", "贡献比"], [
        ["【成员1】", "【学号】", "2023", "【院系】", "【班级】", "统一训练配置、YOLOv8s 基线与结果整理", "【%】"],
        ["【成员2】", "【学号】", "2023", "【院系】", "【班级】", "LUD/FBRT/PRNet 等改进模型训练与横向比较", "【%】"],
        ["【成员3】", "【学号】", "2023", "【院系】", "【班级】", "IF-YOLO 模块、WIoU 与 CSFM-Lite 消融", "【%】"],
        ["【成员4】", "【学号】", "2023", "【院系】", "【班级】", "WFF 稳定性测试、可视化与报告整理", "【%】"],
    ], [1000, 1200, 800, 1400, 1400, 2600, 960], font_size=8)
    heading(doc, "1.1　具体工作说明", 2)
    for text in [
        "完成 VisDrone 原始标注到 YOLO 格式的转换脚本整理，并对数据目录、训练集/验证集/测试集路径进行统一说明。",
        "在第 13 周测试不同输入尺寸、学习率与优化器组合，最后选择 960 输入尺寸、AdamW、lr0=0.003 等统一训练设置。",
        "在统一配置下完成 YOLOv8s、LUD-YOLO、HS-FPN、PRNet、FBRT-YOLO 与 IF-YOLO 的模型训练和横向比较。",
        "围绕 IF-YOLO 继续完成 WIoU v3、CSFM-Lite、P5-MKP-Lite、DySample、LDR 与 WFF 等模块的实现、安装脚本、训练脚本与结果分析。",
        "完成多随机种子稳定性实验与 1280 输入尺寸复杂度检查，整理实验曲线、混淆矩阵、预测可视化和代码附录。",
    ]:
        bullet(doc, text)

    heading(doc, "2　课题概述", 1)
    heading(doc, "2.1　课题背景与任务内容", 2)
    para(doc, "无人机平台能以较低成本获取大范围、俯视角和动态变化的场景图像，因此在交通监控、城市治理、灾害救援、安防巡检和低空经济应用中具有重要价值。与普通地面视角相比，无人机图像中的目标往往面积小、数量多、分布密集，且受到拍摄高度、俯仰角、光照、运动模糊和遮挡的影响，检测难度明显增大。")
    para(doc, "本课题的输入是无人机航拍图像，输出是图像中十类常见目标的位置框、类别和置信度。我们使用 YOLOv8s 作为基础检测器，围绕 VisDrone 数据集开展训练和验证，目标是在尽量保持模型可部署性的前提下提升小目标和密集目标检测能力。")
    table(doc, ["类别编号", "类别名称", "典型难点"], [
        [0, "pedestrian", "目标尺度小，常与人群、阴影混杂"],
        [1, "person", "姿态变化大，遮挡较多"],
        [2, "bicycle", "外观细长，容易与背景纹理混淆"],
        [3, "car", "样本量大，检测相对稳定"],
        [4, "van", "与 car/truck 外观相近"],
        [5, "truck", "尺度跨度大，数量较少"],
        [6, "tricycle", "类别边界模糊，样本较少"],
        [7, "awning-tricycle", "遮挡和外观变化明显"],
        [8, "bus", "目标较大但样本数量有限"],
        [9, "motor", "小目标、密集和运动模糊明显"],
    ], [1000, 2200, 6160], font_size=8.5)
    caption(doc, "表 2-1　VisDrone 检测类别及主要难点")
    for i in range(1, 5):
        add_image(doc, imgs.get(f"sample_{i}", ""), f"图 2-{i}　VisDrone 测试集中典型无人机航拍样例", width=5.6)
    heading(doc, "2.2　国内外研究现状", 2)
    para(doc, "目标检测方法经历了从手工特征与滑动窗口，到两阶段深度检测器，再到单阶段实时检测器的发展。Faster R-CNN 等两阶段方法精度较高，但推理速度和工程部署成本较高；SSD、YOLO 系列等单阶段方法将候选框生成与类别回归合并到一个网络中，更适合实时检测与边缘部署。")
    para(doc, "在无人机视角目标检测中，研究重点通常集中在三类问题：一是小目标特征在下采样过程中容易丢失，需要更强的浅层信息保持机制；二是不同尺度目标同时出现，需要更合理的多尺度特征融合；三是密集目标定位要求更高，需要改进边界框回归损失或后处理策略。IF-YOLO、LUD-YOLO、FBRT-YOLO 等方法均围绕上述问题提出不同结构改进。")
    table(doc, ["方法类别", "代表思路", "优点", "不足"], [
        ["YOLOv8s 基线", "C2f + PAN-FPN + 解耦检测头", "速度快，工程成熟，复现实验方便", "小目标信息易在下采样中衰减"],
        ["LUD-YOLO", "轻量化上采样与细节增强", "严格定位指标较好", "综合召回和 mAP50 不一定最优"],
        ["PRNet", "引入额外采样/高分辨率特征", "关注小目标", "训练成本和显存压力较高"],
        ["IF-YOLO", "IPFA + CSFM + FGAFPN", "Recall 与 mAP50 优势明显", "原始 CSFM 结构参数量较大"],
        ["本项目最终方案", "IF-YOLO + WIoU + CSFM-Lite + WFF", "兼顾定位、轻量融合与稳定性", "仍未解决所有极小目标漏检"],
    ], [1400, 2400, 2600, 2960], font_size=8.3)
    caption(doc, "表 2-2　相关方法比较")
    heading(doc, "2.3　研究意义与应用场景", 2)
    para(doc, "本课题的意义主要体现在三个方面。第一，从算法角度看，它检验了 YOLOv8s 在无人机小目标场景中的瓶颈，并通过消融实验分析不同模块的实际贡献；第二，从工程角度看，项目形成了可复现实验脚本、统一配置、模型检查脚本和结果汇总脚本；第三，从应用角度看，检测结果可为交通流量统计、异常事件发现、目标跟踪与区域巡检提供基础感知能力。")

    heading(doc, "3　总体方案与方法设计", 1)
    heading(doc, "3.1　总体技术路线", 2)
    para(doc, "本项目的整体路线可以概括为“数据准备—统一训练设置—模型横向筛选—IF-YOLO 深入改进—稳定性验证—固定图像展示”。我们没有一开始就直接堆叠模块，而是先用统一训练参数把不同模型放在同一比较框架下，再选择表现最稳的 IF-YOLO 继续迭代。")
    table(doc, ["阶段", "输入", "处理", "输出"], [
        ["数据准备", "VisDrone 原始图像与标注", "标注格式转换、路径整理、类别映射", "YOLO 格式数据集"],
        ["统一配置筛选", "YOLOv8s 基线", "输入尺寸、学习率、优化器和增强策略比较", "后续统一训练配置"],
        ["模型横向比较", "多种 YOLO 改进结构", "同数据、同参数训练与验证", "确定 IF-YOLO 为主线"],
        ["模块改进", "IF-YOLO", "WIoU、CSFM-Lite、MKP、DySample、WFF 等消融", "最终候选模型"],
        ["稳定性验证", "A/B 两组模型", "seed=0/1/2 多次训练", "均值、标准差和最终推荐"],
    ], [1350, 2200, 3500, 2310], font_size=8.5)
    caption(doc, "表 3-1　项目总体技术路线")
    heading(doc, "3.2　数据预处理与增强", 2)
    para(doc, "VisDrone 原始标注以左上角坐标、宽高、类别等字段保存。YOLO 训练要求每个目标表示为归一化后的中心点坐标和宽高，因此我们编写转换脚本，跳过 ignored regions，将类别编号从 VisDrone 的 1 起始调整为 YOLO 的 0 起始。")
    table(doc, ["配置项", "取值", "说明"], [
        ["输入尺寸", "960×960", "第 13 周实验中比 640/800 更适合本数据集"],
        ["batch", "4", "兼顾显存与训练稳定性"],
        ["mosaic", "1.0", "增强密集目标场景泛化能力"],
        ["copy_paste", "0.2", "增加目标组合变化"],
        ["mixup", "0.1", "缓解过拟合并扩展样本分布"],
        ["degrees / scale", "5.0 / 0.3", "模拟拍摄角度与尺度变化"],
        ["seed", "0", "单次对比实验固定随机性"],
    ], [1800, 1900, 5660], font_size=8.5)
    caption(doc, "表 3-2　统一数据增强与训练配置")
    add_image(doc, imgs.get("fig_labels", ""), "图 3-1　训练数据标签分布可视化", width=5.6)
    heading(doc, "3.3　基线模型/基础算法", 2)
    para(doc, "YOLOv8s 是本课程设计的基础网络，具有较好的速度、参数量和成熟工具链。其骨干网络负责多层特征提取，颈部结构完成多尺度特征融合，检测头输出边界框、类别置信度和目标置信度。对于无人机视角任务，YOLOv8s 的优势是训练和部署成本低，劣势是小目标在连续下采样后容易丢失细节。")
    para(doc, "训练目标可以写成定位损失、分类损失和分布式焦点损失的加权组合：")
    para(doc, "L = λbox Lbox + λcls Lcls + λdfl Ldfl", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, "其中 Lbox 约束预测框与真实框的几何重合，Lcls 约束类别判断，Ldfl 用于提升边界框回归精度。本项目统一设置 box=7.5、cls=2.0，并保留 YOLOv8 默认的 DFL 机制。")
    heading(doc, "3.4　改进模型与关键模块", 2)
    heading(doc, "3.4.1　IF-YOLO 主体结构", 3)
    para(doc, "IF-YOLO 的主要思想是从下采样和特征融合两个环节增强无人机小目标检测。IPFA 模块在下采样时尽量保留空间细节，CSFM 模块在多尺度融合时抑制冲突并强化有效响应，FGAFPN 则通过四尺度特征金字塔增加高分辨率检测分支。第 15 周横向实验显示，IF-YOLO 在 Recall 和 mAP50 上优于其他候选模型，因此被确定为后续改进基础。")
    heading(doc, "3.4.2　WIoU v3 定位损失", 3)
    para(doc, "WIoU v3 的引入目标是提升边界框定位质量。普通 IoU 类损失对不同质量样本的关注不够灵活，在密集小目标中容易被异常样本或低质量框干扰。WIoU 通过动态聚焦机制调整样本权重，使训练更关注中等质量、仍有改进空间的边界框。实验中，WIoU 对 mAP50 影响不大，但明显提升 mAP50-95，说明它主要改善严格 IoU 阈值下的定位质量。")
    heading(doc, "3.4.3　CSFM-Lite 轻量特征融合", 3)
    para(doc, "原始 CSFM 能提升多尺度融合能力，但参数量和计算量较大。CSFM-Lite 将较重的冲突抑制和空间融合结构替换为全局池化、ECA 通道加权与深度可分离空间分支，在保持多尺度融合思想的同时显著降低模型规模。该模块是第 16 周实验中最有效的结构改进之一。")
    heading(doc, "3.4.4　WeightedFeatureFusion 最终稳定性模块", 3)
    para(doc, "第 17 周进一步在 IF-YOLO + WIoU v3 + CSFM-Lite 基础上，只替换 Neck 中 P3→P2 的高分辨率融合节点，加入 WeightedFeatureFusion。该模块为输入特征学习可归一化的融合权重，相比直接 Concat 更能控制浅层细节与语义特征的贡献。它只增加 4098 个参数和约 0.5 GFLOPs，却在多随机种子实验中表现出更高的 mAP50-95 均值和更小波动。")
    heading(doc, "3.5　损失函数与优化策略", 2)
    para(doc, "优化器统一采用 AdamW，初始学习率 lr0=0.003，最终学习率比例 lrf=0.01，patience=20。相比直接使用默认参数，统一配置使不同模型之间的比较更公平，也减少了“某个模型只是因为参数更合适而获胜”的可能。")
    heading(doc, "3.6　算法流程与复杂度分析", 2)
    table(doc, ["模型", "Params", "GFLOPs@960", "GFLOPs@1280", "说明"], [
        ["YOLOv8s", "11.139M", "约 64.5*", "114.7", "基础结构，速度快"],
        ["Basic IF-YOLO", "39.438M", "约 161.5*", "287.2", "原始 CSFM 较重"],
        ["IF + WIoU + CSFM-Lite", "11.291M", "105.9", "188.3", "轻量化主干候选"],
        ["Final WFF", "11.295M", "106.4", "189.1", "最终稳定性候选"],
    ], [2100, 1500, 1600, 1600, 2560], font_size=8.5)
    caption(doc, "表 3-3　模型复杂度比较（*为按项目记录近似说明，1280 为第 18 周检查输出）")
    add_image(doc, ASSET_DIR / "main_model_comparison.png", "图 3-2　主要模型 mAP50 与 mAP50-95 对比", width=6.1)

    heading(doc, "4　实验设计与结果分析", 1)
    heading(doc, "4.1　实验环境与统一训练条件", 2)
    para(doc, "实验主要使用 Ultralytics YOLO 框架完成，训练脚本统一写入项目目录，避免在命令行中随意改动关键超参数。部分实验在服务器环境运行，部分结构检查在本地或 course_ai 环境完成。由于不同阶段硬件可能不同，训练时长仅作为记录，不作为结构优劣判断依据。")
    table(doc, ["项目", "配置"], [
        ["训练框架", "Ultralytics YOLOv8，项目中固定训练入口脚本"],
        ["基础权重", "yolov8s.pt"],
        ["训练轮数", "100 epochs"],
        ["输入尺寸", "主要对比 960；第 18 周补充 1280 复杂度与展示包"],
        ["优化器", "AdamW"],
        ["学习率", "lr0=0.003, lrf=0.01"],
        ["评价指标", "Precision、Recall、mAP50、mAP50-95、Params、GFLOPs、FPS"],
    ], [2200, 7160], font_size=8.8)
    heading(doc, "4.2　数据集划分与统计", 2)
    para(doc, "项目使用 VisDrone2019 检测数据，并转换为 YOLO 格式。训练、验证和测试目录均按 images/ 与 labels/ 分开存放，数据集配置文件只需要指向包含 images 和 labels 的根目录。由于数据集体积较大，完整图像、训练权重和 runs 输出不纳入 Git 仓库，只在本地和服务器保留。")
    heading(doc, "4.3　评价指标", 2)
    para(doc, "Precision 表示预测为正的目标中有多少是真正目标；Recall 表示真实目标中有多少被模型检出；mAP50 是 IoU 阈值为 0.5 时的平均精度；mAP50-95 是 IoU 从 0.5 到 0.95 多阈值平均后的结果，对定位质量要求更严格。本项目将 mAP50-95 作为综合指标，同时关注 Recall 和 mAP50，因为无人机场景中漏检会直接影响应用价值。")
    heading(doc, "4.4　对比实验", 2)
    table(doc, ["模型", "最佳轮次", "Precision", "Recall", "mAP50", "mAP50-95", "结论"], [
        ["YOLOv8s", 89, "0.5990", "0.4841", "0.4856", "0.2670", "基础基线"],
        ["LUD-YOLO", 70, "0.5915", "0.4865", "0.4847", "0.2797", "严格定位较好"],
        ["HS-FPN", 57, "0.5842", "0.4787", "0.4791", "0.2613", "未超过基线"],
        ["PRNet", 85, "0.5454", "0.4285", "0.4245", "0.2287", "训练成本较高且效果弱"],
        ["FBRT-YOLO", 85, "0.5995", "0.4814", "0.4863", "0.2636", "Precision 略好但综合不足"],
        ["IF-YOLO", 89, "0.5965", "0.4999", "0.5103", "0.2797", "综合表现最好，选为主线"],
    ], [1650, 1000, 1100, 1100, 1100, 1200, 3220], font_size=8.2)
    caption(doc, "表 4-1　第 15 周六种模型横向比较")
    para(doc, "可以看到，IF-YOLO 的 mAP50 达到 0.5103，Recall 达到 0.4999，是六种模型中最能检出目标的方案；LUD-YOLO 的 mAP50-95 与 IF-YOLO 几乎持平，但 mAP50 和 Recall 低于 IF-YOLO。因此我们没有简单追求单项第一，而是选择综合能力更强、后续改进空间更明确的 IF-YOLO。")
    add_image(doc, imgs.get("fig_yolo_results", ""), "图 4-1　YOLOv8s 基线训练曲线", width=5.9)
    add_image(doc, imgs.get("fig_if_results", ""), "图 4-2　IF-YOLO 训练曲线", width=5.9)
    heading(doc, "4.5　消融实验", 2)
    table(doc, ["模型/模块", "最佳轮次", "Precision", "Recall", "mAP50", "mAP50-95", "相对基础 IF-YOLO"], [
        ["基础 IF-YOLO", 89, "0.59654", "0.49985", "0.51025", "0.27970", "-"],
        ["IF-YOLO + WIoU v3", 84, "0.60004", "0.50065", "0.51042", "0.29014", "+0.01044"],
        ["IF + WIoU + CSFM-Lite", 89, "0.60453", "0.50807", "0.51578", "0.30003", "+0.02033"],
        ["+ P5-MKP-Lite", 88, "0.60144", "0.50032", "0.50919", "0.28740", "+0.00770"],
        ["DySample P4→P3", 88, "0.59180", "0.49164", "0.51019", "0.28734", "低于 CSFM-Lite"],
        ["DySample P3→P2", 90, "0.60072", "0.49174", "0.51421", "0.28799", "低于 CSFM-Lite"],
        ["DySample 双节点", 89, "0.60604", "0.48764", "0.51224", "0.29768", "接近但未超过"],
        ["LDR", 85, "0.61455", "0.49408", "0.51823", "0.29007", "mAP50 高但定位弱"],
        ["WFF 单次", 86, "0.60519", "0.48976", "0.51445", "0.29896", "接近 CSFM-Lite"],
    ], [1750, 900, 1000, 1000, 1000, 1100, 2610], font_size=7.8)
    caption(doc, "表 4-2　IF-YOLO 后续模块消融结果")
    para(doc, "WIoU v3 带来的主要收益体现在 mAP50-95，说明定位质量变好；CSFM-Lite 在保持轻量化的同时进一步提升 Precision、Recall、mAP50 和 mAP50-95，是单次实验中最明显的改进。P5-MKP-Lite 没有形成叠加收益，DySample 单节点效果较弱，双节点和 WFF 接近 CSFM-Lite，但仍需稳定性验证。")
    add_image(doc, imgs.get("fig_final_results", ""), "图 4-3　IF-YOLO + WIoU v3 + CSFM-Lite 训练曲线", width=5.9)
    add_image(doc, imgs.get("fig_final_pr", ""), "图 4-4　CSFM-Lite 最优模型 PR 曲线", width=5.6)
    add_image(doc, imgs.get("fig_final_conf", ""), "图 4-5　CSFM-Lite 归一化混淆矩阵", width=5.4)
    heading(doc, "4.6　稳定性测试与最终模型选择", 2)
    para(doc, "单次训练的最高值可能受到随机初始化、数据增强采样和训练过程波动影响。第 17 周我们将 A 组设为 IF-YOLO + WIoU v3 + CSFM-Lite，将 B 组设为在 A 组基础上加入 WFF(P3→P2)，分别使用 seed=0/1/2 进行训练。")
    table(doc, ["组别", "模型", "mAP50-95 均值", "标准差", "mAP50 均值", "Recall 均值", "结论"], [
        ["A", "IF + WIoU + CSFM-Lite", "0.28613", "0.00354", "0.51113", "0.50220", "单次最好但波动较大"],
        ["B", "A + WFF(P3→P2)", "0.29380", "0.00138", "0.51180", "0.51070", "均值更高且更稳定"],
    ], [700, 2600, 1300, 1000, 1100, 1100, 1560], font_size=8.3)
    caption(doc, "表 4-3　三随机种子稳定性结果")
    para(doc, "从稳定性角度看，B 组 mAP50-95 均值比 A 组高 0.00767，标准差更小，Recall 也更高。因此最终展示模型选择 B 组，即 IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3→P2)。同时，报告中保留 CSFM-Lite 单次最高结果，用于说明模型在最优单次训练中的上限。")
    add_image(doc, ASSET_DIR / "multiseed_stability.png", "图 4-6　三随机种子稳定性对比", width=5.8)
    add_image(doc, imgs.get("fig_wff_results", ""), "图 4-7　最终 WFF 候选模型训练曲线", width=5.9)
    add_image(doc, imgs.get("fig_wff_pr", ""), "图 4-8　最终 WFF 候选模型 PR 曲线", width=5.6)
    heading(doc, "4.7　可视化结果与误差分析", 2)
    para(doc, "从预测图可以看出，模型对车辆、公交车等尺寸较大且纹理清晰的类别检测效果较好；对行人、person、bicycle、awning-tricycle 等小目标和细长目标仍存在漏检。主要原因包括：目标在图像中占比过小；密集区域中目标之间互相遮挡；部分类别外观差异小；YOLO txt 验证流程没有直接输出 COCO 风格 APsmall，因此对极小目标的定量分析还不够细。")
    add_image(doc, imgs.get("fig_final_val", ""), "图 4-9　CSFM-Lite 候选模型验证集预测示例", width=6.1)
    add_image(doc, imgs.get("fig_wff_val", ""), "图 4-10　最终 WFF 候选模型验证集预测示例", width=6.1)
    heading(doc, "4.8　阶段性结果图版补充", 2)
    para(doc, "为了便于复核完整实验过程，本节集中放置部分代表性训练曲线、PR 曲线、混淆矩阵和预测结果图。正文分析只选取关键图，图版补充则保留更多阶段性证据，便于后续答辩时按需要查找。")
    gallery = [
        (ROOT / "Model/runs/detect/week15_yolov8s_baseline/BoxPR_curve.png", "图 4-11　YOLOv8s 基线 PR 曲线"),
        (ROOT / "Model/runs/detect/week15_yolov8s_baseline/confusion_matrix_normalized.png", "图 4-12　YOLOv8s 基线归一化混淆矩阵"),
        (ROOT / "Model/runs/detect/week15_lud_yolo/results.png", "图 4-13　LUD-YOLO 训练曲线"),
        (ROOT / "Model/runs/detect/week15_lud_yolo/BoxPR_curve.png", "图 4-14　LUD-YOLO PR 曲线"),
        (ROOT / "本周进展说明及下周计划报告/十五周/train-2/results.png", "图 4-15　HS-FPN 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十五周/train-35/results.png", "图 4-16　PRNet 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十五周/fbrt_yolov8s_visdrone_pretrained_week18-4/results.png", "图 4-17　FBRT-YOLO 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/ablation_03_csfm/ablation_03_csfm/results.png", "图 4-18　CSFM 消融训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/detect/ablation_04_fgafpn/results.png", "图 4-19　FGAFPN 消融训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/IPFA_new/ablation_02_ipfa_stable/ablation_02_ipfa_stable/results.png", "图 4-20　IPFA 稳定版消融训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou-2/week15_if_yolov8s_wiou-2/results.png", "图 4-21　IF-YOLO + WIoU v3 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_result/week15_if_yolov8s_wiou_csfm_lite_mkp/week15_if_yolov8s_wiou_csfm_lite_mkp/results.png", "图 4-22　P5-MKP-Lite 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_p4p3/results.png", "图 4-23　DySample P4→P3 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_p3p2/results.png", "图 4-24　DySample P3→P2 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_p4p3_p3p2/results.png", "图 4-25　DySample 双节点训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_ldr/results.png", "图 4-26　LocalDetailRefine 训练曲线"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/confusion_matrix_normalized.png", "图 4-27　WFF 归一化混淆矩阵"),
        (ROOT / "本周进展说明及下周计划报告/十七周/results/新建文件夹/新建文件夹/week15_if_wiou_csfmlite_dys_best_wff-2/val_batch1_pred.jpg", "图 4-28　WFF 验证集预测示例 2"),
    ]
    for path, title in gallery:
        add_image_plate(doc, path, title, width=6.15)

    heading(doc, "5　系统实现与应用展示", 1)
    heading(doc, "5.1　系统总体架构", 2)
    para(doc, "项目实现并不是单一训练脚本，而是一套可复现实验流程：Dataset 目录负责数据转换；Model 目录负责基线与通用训练脚本；各周报告目录保留阶段性模型包、模块安装脚本、训练脚本、结果汇总脚本和可视化图；Reference 目录保存相关论文；report 目录保存课程报告模板和最终报告。")
    table(doc, ["目录/文件", "作用"], [
        ["Dataset/visdrone2yolo.py", "将 VisDrone 原始标注转换为 YOLO 标签"],
        ["Model/train_visdrone.py", "YOLOv8s 基线训练入口"],
        ["Model/train_config.yaml", "第 15 周后统一训练配置"],
        ["十六周/week15_final_training_package", "IF-YOLO、WIoU、CSFM-Lite、MKP 训练包"],
        ["十七周/week17", "WFF 与多随机种子稳定性实验包"],
        ["18周", "1280 输入尺寸复杂度检查与固定测试图展示包"],
    ], [3200, 6160], font_size=8.5)
    heading(doc, "5.2　功能模块设计", 2)
    para(doc, "核心实现模块包括数据转换、模型注册、训练入口、模型检查、结果汇总和固定图片预测。自定义模块需要复制到 Ultralytics 对应模块目录，并在 tasks.py 的模型解析过程中注册通道推导逻辑。为了减少人工改源码出错，项目将这些步骤写成 install_weekXX_modules.py，保证训练环境可重复搭建。")
    heading(doc, "5.3　系统运行流程", 2)
    for text in [
        "准备 VisDroneYOLO 数据集，并在 visdrone.yaml 中填写真实 path。",
        "安装 Python、PyTorch、Ultralytics，并确认 GPU 可用。",
        "运行自定义模块安装脚本，检查 IPFA、CSFM-Lite、WFF 等模块能否构建。",
        "使用 train_experiment.py 或 train_1280_experiment.py 按模型名称启动训练。",
        "训练结束后读取 results.csv、results.png、PR 曲线、混淆矩阵和预测图。",
        "运行 summarize 脚本汇总多随机种子或多模型结果。",
    ]:
        bullet(doc, text)
    heading(doc, "5.4　应用效果展示", 2)
    para(doc, "第 18 周准备了固定 15 张 VisDrone test-challenge 图片，用于让不同模型在完全相同的图像上输出预测结果。该展示包不参与训练和验证指标计算，主要用于课程答辩中直观展示模型对车辆、行人、摩托车等目标的检测效果。")

    heading(doc, "6　结束语", 1)
    heading(doc, "6.1　工作总结", 2)
    para(doc, "本课程设计从 YOLOv8s 基线出发，围绕无人机视角目标检测完成了比较完整的实验闭环。我们没有只跑一个最终模型，而是先确定统一训练配置，再横向比较多个改进方向，最后围绕 IF-YOLO 做进一步模块消融和稳定性验证。最终选择 IF-YOLO + WIoU v3 + CSFM-Lite + WFF(P3→P2)，主要原因是它在三随机种子测试中表现更稳定，且复杂度仅比无 WFF 版本略有增加。")
    heading(doc, "6.2　不足与展望", 2)
    para(doc, "本项目仍有不足。第一，虽然 mAP50-95 有提升，但小目标和严重遮挡类别仍然是短板；第二，当前验证流程主要使用 YOLO 自带指标，未进一步启用 COCO-json 评估以获得 APsmall、APmedium 等尺度指标；第三，部分模型训练硬件不同，训练时长只能参考，不能直接比较结构效率；第四，最终模型尚未部署到真实无人机视频流或边缘设备上。后续可以继续从尺度敏感评价、轻量化部署、视频时序信息和更强的数据增强策略四个方向改进。")

    heading(doc, "参考文献", 1)
    refs = [
        "[1] Ultralytics. YOLOv8 documentation and implementation.",
        "[2] Du D., Zhu P., Wen L., et al. VisDrone-DET2019: The Vision Meets Drone Object Detection in Image Challenge.",
        "[3] IF-YOLO: improved YOLO-based object detection method for drone-view or small-object detection scenarios.",
        "[4] LUD-YOLO: lightweight UAV/drone object detection method.",
        "[5] FBRT-YOLO: feature-enhanced YOLO variant used as a comparative method in this course project.",
        "[6] Zheng Z., Wang P., Liu W., et al. Distance-IoU Loss: Faster and Better Learning for Bounding Box Regression.",
        "[7] 项目本地实验记录：D:\\大三下\\模式识别课设\\findings.md、progress.md 与各周阶段报告。",
    ]
    for r in refs:
        para(doc, r)

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "附录 A　核心代码与配置文件", 1)
    heading(doc, "A.1　项目目录结构", 2)
    add_code_block(doc, """.
├─ Dataset/                         数据转换脚本与数据说明
├─ Model/                           YOLOv8s 基线、统一配置和训练脚本
├─ Reference/                       参考论文
├─ report/                          模板与最终报告
├─ tools/                           报告生成、文档提取、结果对比工具
└─ 本周进展说明及下周计划报告/       第 12-18 周阶段材料、实验包和结果""", 20)
    heading(doc, "A.2　VisDrone 标注转换代码", 2)
    add_code_block(doc, read_text(ROOT / "Dataset/visdrone2yolo.py"), 65)
    doc.add_page_break()
    heading(doc, "A.3　统一训练配置", 2)
    add_code_block(doc, read_text(ROOT / "Model/train_config.yaml"), 60)
    heading(doc, "A.3.1　VisDrone 数据集配置", 3)
    add_code_block(doc, read_text(ROOT / "Model/visdrone.yaml"), 60)
    doc.add_page_break()
    heading(doc, "A.4　YOLOv8s 基线训练入口", 2)
    add_code_block(doc, read_text(ROOT / "Model/train_visdrone.py"), 70)
    heading(doc, "A.4.1　第十五周统一训练入口", 3)
    p = ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_package/train_week15_experiment.py"
    if p.exists():
        add_code_block(doc, read_text(p), 120)
    doc.add_page_break()
    heading(doc, "A.5　IF-YOLO 模块安装与训练说明摘要", 2)
    p = ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_package/README_训练说明.md"
    add_code_block(doc, read_text(p), 70)
    heading(doc, "A.5.1　IF-YOLO + WIoU + CSFM-Lite 配置片段", 3)
    p = ROOT / "本周进展说明及下周计划报告/十六周/week15_final_training_package/if_yolov8s_wiou_csfm_lite.yaml"
    if p.exists():
        add_code_block(doc, read_text(p), 120)
    doc.add_page_break()
    heading(doc, "A.6　WeightedFeatureFusion 核心模块", 2)
    p = ROOT / "本周进展说明及下周计划报告/十七周/week17/modules/weighted_fusion.py"
    if p.exists():
        add_code_block(doc, read_text(p), 140)
    else:
        add_code_block(doc, "WeightedFeatureFusion 模块位于第十七周 week17/modules/weighted_fusion.py，当前路径未找到。", 20)
    heading(doc, "A.6.1　最终 WFF YAML 配置", 3)
    p = ROOT / "本周进展说明及下周计划报告/十七周/week17/configs/if_yolov8s_wiou_csfmlite_wff_p3p2.yaml"
    if p.exists():
        add_code_block(doc, read_text(p), 120)
    doc.add_page_break()
    heading(doc, "A.7　稳定性实验训练入口摘要", 2)
    p = ROOT / "本周进展说明及下周计划报告/十七周/week17/scripts/train_experiment.py"
    if p.exists():
        add_code_block(doc, read_text(p), 130)
    else:
        add_code_block(doc, "稳定性实验训练入口位于第十七周 week17/scripts/train_experiment.py，当前路径未找到。", 20)
    heading(doc, "A.8　第 18 周 1280 输入检查脚本摘要", 2)
    p = ROOT / "本周进展说明及下周计划报告/18周/scripts/check_models_1280.py"
    if p.exists():
        add_code_block(doc, read_text(p), 130)
    heading(doc, "A.9　固定图片预测脚本摘要", 2)
    p = ROOT / "本周进展说明及下周计划报告/18周/scripts/predict_15_images.py"
    if p.exists():
        add_code_block(doc, read_text(p), 130)

    for section in doc.sections:
        add_page_number(section)
    doc.core_properties.title = "基于 YOLOv8s 与 IF-YOLO 改进的无人机视角目标检测课程设计报告"
    doc.core_properties.author = "第41组"
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_report())
