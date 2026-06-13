from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"D:\大三下\模式识别课设\本周进展说明及下周计划报告\十四周")
OUT_FILE = OUT_DIR / "41-无人机视角下目标检测-十三周工作进展及下周计划报告.docx"


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def set_east_asia_font(run, font_name):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFC7D1"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_cell_margins(table)

    header_row = table.rows[0]
    for idx, text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = text
        shade_cell(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    for values in rows:
        row = table.add_row()
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.text = str(text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    return table


def para(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D7DEE8")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FC")
    p = cell.paragraphs[0]
    for line in code.strip().splitlines():
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.add_run("\n")
    return table


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("小组本周工作总结以及下周安排")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    set_east_asia_font(run, "SimHei")
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("无人机视角下目标检测 - 第十四周")
    srun.font.size = Pt(11)
    srun.font.color.rgb = RGBColor.from_string("555555")
    set_east_asia_font(srun, "SimHei")
    subtitle.paragraph_format.space_after = Pt(12)

    doc.add_heading("一、本周（十四周）工作进展", level=1)
    para(
        doc,
        "本周我们小组在上一周统一 YOLOv8s 训练参数的基础上，重点从网络结构改进入手，围绕无人机航拍图像中小目标密集、"
        "尺度变化大、遮挡多、类别不均衡等问题，分别调研并尝试了多种基于 YOLOv8s 的改进方案。整体来看，本周工作从单纯调参"
        "推进到网络结构优化阶段，主要完成了四个方向的方案整理、模型修改和训练结果分析。"
    )

    bullet(
        doc,
        "完成原始 YOLOv8s 与改进模型的效果对比。原始 YOLOv8s 在验证集上 Precision 为 0.597、Recall 为 0.470、"
        "mAP50 为 0.476、mAP50-95 为 0.288，为后续结构改进提供了统一基线。"
    )
    bullet(
        doc,
        "完成 LUD-YOLOv8s 结构改进实验。该方案在 YOLOv8s 中引入 C2fBRA 注意力模块和 ASFF3 自适应特征融合模块，"
        "训练 100 epoch 后 mAP50 达到 0.488，Recall 提升到 0.486，说明模型检出能力有所增强。"
    )
    bullet(
        doc,
        "整理 FBRT-YOLO 改进思路。该方案参考 AAAI 2025 的航拍目标检测网络，重点引入 FCM 特征互补映射模块和 MKP 多核感知单元，"
        "用于缓解下采样过程中的空间信息丢失，并增强小目标的多尺度上下文建模能力。"
    )
    bullet(
        doc,
        "整理 PRNet 改进思路。该方案主要增加 ESSamp 下采样模块和 PRN 特征复用模块，强调在下采样阶段保留细节，并在颈部多次复用"
        "P2、P3 等高分辨率浅层特征，以提升小目标定位能力。"
    )
    bullet(
        doc,
        "整理 IF-YOLO 改进思路。该方案引入 IPFA、CSFM 和 FGAFPN 三个模块，分别用于保留下采样信息、抑制融合冲突信息、增强细粒度"
        "多尺度特征交互；实验结果中改进 YOLOv8s 的 mAP50 达到 0.490，mAP50-95 达到 0.288。"
    )

    doc.add_heading("本周各方案汇总", level=2)
    add_table(
        doc,
        ["方向", "主要改进模块", "改进目的", "当前结果或进展"],
        [
            [
                "LUD-YOLOv8s",
                "C2fBRA、ASFF3",
                "增强注意力表达和多尺度自适应融合能力",
                "P=0.601，R=0.486，mAP50=0.488，mAP50-95=0.280；mAP50 和召回率较原始 YOLOv8s 小幅提升。",
            ],
            [
                "FBRT-YOLO",
                "FCM、MKP",
                "解决浅层空间信息丢失和小目标跨尺度感知不足问题",
                "完成方案调研和训练总结，180 epoch 训练约 13.6 小时，mAP50 约 0.498，训练速度较快。",
            ],
            [
                "PRNet",
                "ESSamp、PRN",
                "在下采样时保留细节，并在特征融合阶段重复利用浅层高分辨率特征",
                "完成模块原理整理，后续需要进一步统一训练并补充量化对比。",
            ],
            [
                "IF-YOLO",
                "IPFA、CSFM、FGAFPN",
                "保留下采样信息、抑制融合冲突、加强细粒度特征金字塔融合",
                "改进模型 P=0.595，R=0.481，mAP50=0.490，mAP50-95=0.288，相比基线 mAP50 提升 0.8%。",
            ],
        ],
        [1600, 2100, 2700, 2960],
    )

    para(
        doc,
        "从本周结果看，结构改进确实能在部分指标上带来提升，尤其是 mAP50、Recall 或小目标类别上出现一定改善。"
        "但各方案的训练轮数、实现细节和评价方式还没有完全统一，因此目前更适合作为结构改进方向的阶段性探索，"
        "还不能直接作为最终结论。"
    )

    doc.add_heading("本周存在的问题", level=2)
    bullet(
        doc,
        "不同成员采用的改进模块和训练轮数不同，例如 LUD-YOLOv8s 训练 100 epoch，而 FBRT-YOLO 训练 180 epoch，横向对比仍不完全公平。"
    )
    bullet(
        doc,
        "部分类别仍然检测困难，尤其是 bicycle、tricycle、awning-tricycle 等小目标或类别样本较少的目标，mAP 明显低于 car、bus 等大目标类别。"
    )
    bullet(
        doc,
        "部分改进模型虽然提升了 mAP50 或 Recall，但参数量和计算量也明显增加，需要继续评估精度提升是否值得额外计算开销。"
    )
    bullet(
        doc,
        "目前缺少统一消融实验，尚不能明确每个模块对最终指标提升的具体贡献。"
    )

    doc.add_heading("二、下周工作计划", level=1)
    bullet(
        doc,
        "统一训练设置。下周优先统一数据集配置、输入尺寸、batch size、优化器、学习率、epoch、随机种子和评估方式，保证所有结构改进方案可公平对比。"
    )
    bullet(
        doc,
        "完成消融实验。重点对 LUD-YOLOv8s 中的 C2fBRA 和 ASFF3 分别进行单独训练，比较 YOLOv8s、YOLOv8s+C2fBRA、YOLOv8s+ASFF3、"
        "YOLOv8s+C2fBRA+ASFF3 四组结果。"
    )
    bullet(
        doc,
        "补全 PRNet 和 FBRT-YOLO 的统一结果。对已经整理的 PRNet、FBRT-YOLO 方案继续进行复现，补充 mAP50、mAP50-95、参数量、GFLOPs、"
        "单张推理耗时等指标。"
    )
    bullet(
        doc,
        "重点分析小目标类别。针对 bicycle、tricycle、awning-tricycle、person 等弱类别，统计样本数量、误检和漏检情况，结合预测图分析问题来源。"
    )
    bullet(
        doc,
        "整理最终实验材料。将原始 YOLOv8s、各结构改进模型、训练曲线、混淆矩阵、PR 曲线和可视化检测结果统一整理，为最终课设报告做准备。"
    )

    doc.add_heading("三、下周统一训练与对比安排", level=1)
    para(
        doc,
        "下周所有成员优先使用统一参数进行训练和验证，保证不同模型之间只有网络结构不同，其他训练条件尽可能保持一致。"
        "统一训练配置建议如下："
    )
    add_code_block(
        doc,
        """
from ultralytics import YOLO
import multiprocessing

multiprocessing.freeze_support()

if __name__ == '__main__':
    model = YOLO('改进模型对应的yaml或pt文件')
    model.train(
        data='VisDrone.yaml',
        epochs=100,
        imgsz=960,
        batch=4,
        optimizer='AdamW',
        lr0=0.003,
        lrf=0.01,
        patience=20,
        mosaic=1.0,
        copy_paste=0.2,
        mixup=0.1,
        degrees=5.0,
        scale=0.3,
        cls=2.0,
        box=7.5,
        iou=0.5,
        device=0,
        workers=0,
        cache=False,
        amp=True,
        val=True,
        plots=True,
        seed=0
    )
    metrics = model.val()
        """,
    )

    para(
        doc,
        "最终对比时主要统计 Precision、Recall、mAP50、mAP50-95、参数量、GFLOPs、训练耗时和推理速度，并结合分类别结果判断模型是否真正提升了无人机航拍小目标检测能力。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frun = footer.add_run("无人机视角下目标检测 - 第十四周工作进展")
    frun.font.size = Pt(9)
    frun.font.color.rgb = RGBColor.from_string("777777")

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
